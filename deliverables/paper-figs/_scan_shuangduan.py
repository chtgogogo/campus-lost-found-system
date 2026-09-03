# -*- coding: utf-8 -*-
"""扫描毕业论文 docx 中的“双端”类残留措辞与“双码交叉验证”出现次数。

用途：任务 A 修改前后的基线/验证扫描，不修改任何文件。
"""
from __future__ import annotations

import io
import sys
from typing import Dict, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DOCX_PATH: str = (
    r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
)

# 需要统计的关键词（“双端”类残留 + 目标术语）
KEYWORDS: List[str] = [
    "双端验证",
    "双端交叉验证",
    "双端均确认",
    "双端扫码",
    "双端确认",
    "双端",          # 兜底：任何“双端”出现
    "双码交叉验证",
]


def iter_body_paragraphs(doc: DocxDocument) -> List[Tuple[str, Paragraph]]:
    """返回正文段落列表，标签形如 'para#12'。"""
    return [(f"para#{i}", p) for i, p in enumerate(doc.paragraphs)]


def iter_table_paragraphs(doc: DocxDocument) -> List[Tuple[str, Paragraph]]:
    """返回所有表格单元格内的段落，标签形如 'table#3/r2/c1/p0'。"""
    out: List[Tuple[str, Paragraph]] = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    out.append((f"table#{ti}/r{ri}/c{ci}/p{pi}", p))
    return out


def scan(doc: DocxDocument) -> Tuple[Dict[str, int], List[Tuple[str, str]]]:
    """统计关键词出现次数，并返回命中“双端”的位置与全文。"""
    counts: Dict[str, int] = {k: 0 for k in KEYWORDS}
    hits: List[Tuple[str, str]] = []

    for label, p in iter_body_paragraphs(doc) + iter_table_paragraphs(doc):
        text = p.text or ""
        if not text.strip():
            continue
        for kw in KEYWORDS:
            counts[kw] += text.count(kw)
        if "双端" in text:
            hits.append((label, text))
    return counts, hits


def main() -> int:
    doc: DocxDocument = Document(DOCX_PATH)
    counts, hits = scan(doc)

    print("=" * 78)
    print("关键词计数（全文：正文段落 + 表格单元格）")
    print("=" * 78)
    for kw in KEYWORDS:
        print(f"  {kw:<12} = {counts[kw]}")

    print()
    print("=" * 78)
    print(f"含“双端”的位置共 {len(hits)} 处")
    print("=" * 78)
    for label, text in hits:
        print(f"\n--- [{label}] ---")
        print(text)

    # 额外：输出 runs 结构，便于判断能否整段替换
    print()
    print("=" * 78)
    print("命中段落的 runs 明细（用于安全替换）")
    print("=" * 78)
    all_paras = dict(iter_body_paragraphs(doc) + iter_table_paragraphs(doc))
    for label, _ in hits:
        p = all_paras[label]
        print(f"\n--- [{label}] runs={len(p.runs)} ---")
        for i, r in enumerate(p.runs):
            if "双端" in r.text or "双" in r.text or "端" in r.text:
                print(f"   run[{i}] = {r.text!r}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
