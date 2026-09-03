# -*- coding: utf-8 -*-
"""扫描毕业论文 docx：(1) T13 所在表格行上下文；(2) 用例图相关章节原文。

只读脚本，不修改任何文件。
"""
from __future__ import annotations

import io
import re
import sys
from typing import List

from docx import Document
from docx.document import Document as DocxDocument

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DOCX_PATH: str = (
    r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
)

USECASE_PAT: re.Pattern = re.compile(
    r"用例|参与者|actor|游客|未登录|浏览公示|公示栏|角色|失主|拾得者|管理员"
)


def dump_table_row(doc: DocxDocument, table_idx: int, row_idx: int) -> None:
    """打印指定表格指定行（含表头）的全部单元格文本。"""
    table = doc.tables[table_idx]
    print("=" * 78)
    print(f"table#{table_idx}  rows={len(table.rows)}  cols={len(table.columns)}")
    print("=" * 78)
    header = table.rows[0]
    print("[表头] " + " | ".join(c.text.strip() for c in header.cells))
    for ri in (row_idx - 1, row_idx, row_idx + 1):
        if 0 <= ri < len(table.rows):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            print(f"[r{ri}] " + " | ".join(cells))


def dump_usecase_section(doc: DocxDocument) -> None:
    """定位并打印“用例图”相关章节的正文段落。"""
    paras: List[str] = [p.text for p in doc.paragraphs]

    # 1) 找出所有提到“用例图”“图 3.4”的段落索引
    anchors: List[int] = [
        i for i, t in enumerate(paras)
        if ("用例图" in t) or ("图 3.4" in t) or ("图3.4" in t)
    ]
    print()
    print("=" * 78)
    print(f"提到“用例图 / 图3.4”的段落索引: {anchors}")
    print("=" * 78)

    shown: set = set()
    for a in anchors:
        lo, hi = max(0, a - 12), min(len(paras), a + 16)
        for i in range(lo, hi):
            if i in shown:
                continue
            shown.add(i)
            t = paras[i].strip()
            if t:
                print(f"[para#{i}] {t}")
        print("-" * 78)


def dump_role_paragraphs(doc: DocxDocument) -> None:
    """打印所有含角色/用例关键词的段落，便于提取用例名。"""
    print()
    print("=" * 78)
    print("含角色/用例关键词的段落（全文）")
    print("=" * 78)
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        if USECASE_PAT.search(t) and len(t) < 600:
            print(f"[para#{i}] {t}")


def dump_all_tables_brief(doc: DocxDocument) -> None:
    """列出全部表格的首行，便于识别用例表 / 角色表。"""
    print()
    print("=" * 78)
    print("全部表格首行速览")
    print("=" * 78)
    for ti, table in enumerate(doc.tables):
        try:
            head = " | ".join(c.text.strip() for c in table.rows[0].cells)
        except IndexError:
            head = "(空表)"
        print(f"table#{ti} rows={len(table.rows)} :: {head[:150]}")


def main() -> int:
    doc: DocxDocument = Document(DOCX_PATH)
    dump_table_row(doc, table_idx=4, row_idx=13)
    dump_all_tables_brief(doc)
    dump_usecase_section(doc)
    dump_role_paragraphs(doc)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
