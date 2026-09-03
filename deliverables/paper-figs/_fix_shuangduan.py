# -*- coding: utf-8 -*-
"""任务 A：统一毕业论文 docx 中 5 处“双端”类残留措辞 → “双码交叉验证”。

改动清单（严格限定，只动这 5 处）：
  1. para#139            “交接码生成与双端验证”      → “交接码生成与双码交叉验证”
  2. para#217            “双端交叉验证均通过后”      → “双码交叉验证均通过后”
  3. para#299            “生成交接码→双端验证→”      → “生成交接码→双码交叉验证→”
  4. table#4/r13/c1(T13) “双端验证交接”              → “双码交叉验证交接”
  5. table#4/r13/c3(T13) “双端均确认”                → “双码交叉验证均通过”

实现要点：
  - 在 run 级别做替换以完整保留字体/字号等格式；
  - 若目标串跨 run 分裂，则回退为“首 run 承载全文、其余 run 清空”的合并策略；
  - 改前自动备份，改后做全局断言校验（双端类=0，双码交叉验证=基线+5）。
"""
from __future__ import annotations

import io
import shutil
import sys
from datetime import datetime
from typing import Dict, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DOCX_PATH: str = (
    r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
)

# 替换规则，顺序敏感：长串/特例优先
RULES: List[Tuple[str, str]] = [
    ("双端交叉验证", "双码交叉验证"),
    ("双端均确认", "双码交叉验证均通过"),
    ("双端验证", "双码交叉验证"),
]

# 修改后必须全部归零的“双端”类词
FORBIDDEN: List[str] = [
    "双端验证",
    "双端交叉验证",
    "双端均确认",
    "双端扫码",
    "双端确认",
    "双端",
]

TARGET_TERM: str = "双码交叉验证"
EXPECTED_DELTA: int = 5


def collect_paragraphs(doc: DocxDocument) -> List[Tuple[str, Paragraph]]:
    """收集正文段落 + 表格单元格段落（表格单元格按 _tc 去重，避免合并单元格重复计数）。"""
    items: List[Tuple[str, Paragraph]] = [
        (f"para#{i}", p) for i, p in enumerate(doc.paragraphs)
    ]
    seen_tc: set = set()
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                key = id(cell._tc)
                if key in seen_tc:
                    continue
                seen_tc.add(key)
                for pi, p in enumerate(cell.paragraphs):
                    items.append((f"table#{ti}/r{ri}/c{ci}/p{pi}", p))
    return items


def count_terms(doc: DocxDocument, terms: List[str]) -> Dict[str, int]:
    """统计给定词在全文（正文 + 表格）中的出现次数。"""
    counts: Dict[str, int] = {t: 0 for t in terms}
    for _, p in collect_paragraphs(doc):
        text = p.text or ""
        for t in terms:
            counts[t] += text.count(t)
    return counts


def replace_in_paragraph(p: Paragraph, rules: List[Tuple[str, str]]) -> int:
    """在单个段落内按规则替换，返回替换次数。优先 run 级替换以保留格式。"""
    original: str = p.text or ""
    if not any(old in original for old, _ in rules):
        return 0

    n_replaced: int = 0

    # 策略 1：目标串完整落在某个 run 内 —— 直接改该 run，格式零损失
    for run in p.runs:
        run_text: str = run.text or ""
        new_text: str = run_text
        for old, new in rules:
            if old in new_text:
                n_replaced += new_text.count(old)
                new_text = new_text.replace(old, new)
        if new_text != run_text:
            run.text = new_text

    if n_replaced:
        return n_replaced

    # 策略 2：目标串被拆散到多个 run —— 合并到首 run，其余 run 置空
    merged: str = original
    for old, new in rules:
        if old in merged:
            n_replaced += merged.count(old)
            merged = merged.replace(old, new)
    if n_replaced and p.runs:
        p.runs[0].text = merged
        for run in p.runs[1:]:
            run.text = ""
    return n_replaced


def main() -> int:
    # ---------- 1. 基线扫描 ----------
    doc: DocxDocument = Document(DOCX_PATH)
    before: Dict[str, int] = count_terms(doc, FORBIDDEN + [TARGET_TERM])
    baseline_target: int = before[TARGET_TERM]

    print("=" * 78)
    print("【修改前基线】")
    print("=" * 78)
    for k in FORBIDDEN + [TARGET_TERM]:
        print(f"  {k:<12} = {before[k]}")

    # ---------- 2. 备份 ----------
    stamp: str = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path: str = f"{DOCX_PATH}.bak.detailfix2.{stamp}"
    shutil.copy2(DOCX_PATH, backup_path)
    print(f"\n[备份] {backup_path}")

    # ---------- 3. 执行替换 ----------
    print("\n" + "=" * 78)
    print("【执行替换】")
    print("=" * 78)
    total: int = 0
    for label, p in collect_paragraphs(doc):
        old_text: str = p.text or ""
        n: int = replace_in_paragraph(p, RULES)
        if n:
            total += n
            print(f"\n[{label}] 替换 {n} 处")
            print(f"  改前: ...{_ctx(old_text)}")
            print(f"  改后: ...{_ctx(p.text)}")
    print(f"\n合计替换 {total} 处（预期 {EXPECTED_DELTA} 处）")

    # ---------- 4. 保存 ----------
    doc.save(DOCX_PATH)
    print(f"\n[已保存] {DOCX_PATH}")

    # ---------- 5. 重新打开做独立校验 ----------
    doc2: DocxDocument = Document(DOCX_PATH)
    after: Dict[str, int] = count_terms(doc2, FORBIDDEN + [TARGET_TERM])

    print("\n" + "=" * 78)
    print("【修改后校验（重新读盘）】")
    print("=" * 78)
    ok: bool = True
    for k in FORBIDDEN:
        flag = "OK" if after[k] == 0 else "FAIL"
        if after[k] != 0:
            ok = False
        print(f"  {k:<12} = {after[k]}   [{flag}]")

    delta: int = after[TARGET_TERM] - baseline_target
    d_flag = "OK" if delta == EXPECTED_DELTA else "FAIL"
    if delta != EXPECTED_DELTA:
        ok = False
    print(
        f"  {TARGET_TERM:<12} = {after[TARGET_TERM]} "
        f"(基线 {baseline_target} + {delta})   [{d_flag}]"
    )

    print("\n" + "=" * 78)
    print(f"任务A 结论: {'PASS' if ok else 'FAIL'}")
    print("=" * 78)
    return 0 if ok else 1


def _ctx(text: str, width: int = 90) -> str:
    """截取含关键词的上下文片段，便于人工核对。"""
    for kw in ("双码交叉验证", "双端"):
        idx = text.find(kw)
        if idx >= 0:
            lo = max(0, idx - 30)
            return text[lo: lo + width].replace("\n", " ")
    return text[:width]


if __name__ == "__main__":
    raise SystemExit(main())
