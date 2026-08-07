# -*- coding: utf-8 -*-
"""P5 彻底反查：全文档（正文+表格+页眉页脚）搜旧词，并核验 11类+其他 枚举 vs 代码 12 类。"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DOCX = Path(r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx")

doc = Document(str(DOCX))

# 收集所有文本：正文段落 + 表格单元格 + 页眉页脚
texts = []  # (location, text)
for i, p in enumerate(doc.paragraphs):
    texts.append((f"para[{i}]", p.text or ""))
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            texts.append((f"table[{ti}]r{ri}c{ci}", cell.text or ""))
# 页眉页脚
for section in doc.sections:
    for hdr in [section.header, section.first_page_header, section.even_page_header]:
        for p in hdr.paragraphs:
            texts.append(("header", p.text or ""))
    for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
        for p in ftr.paragraphs:
            texts.append(("footer", p.text or ""))

OLD = ["背包", "本子", "玩偶", "娃娃", "doll", "Doll", "DOLL"]
print("=== 旧词全文档反查 ===")
found = False
for loc, t in texts:
    for o in OLD:
        if o in t:
            print(f"  {loc}: ...{t[max(0,t.find(o)-15):t.find(o)+15]}...")
            found = True
if not found:
    print("  ✅ 全文（正文/表格/页眉页脚）搜不到 背包/本子/玩偶/娃娃/doll")

# 核验枚举：找包含 11 个具体类目列举的段落
print("\n=== 类目枚举段核对（应含：手机 钱包 钥匙 书包 行李箱 笔记本电脑 校园卡 眼镜 笔记本 雨伞 水杯 + 其他）===")
CANON = ["手机","钱包","钥匙","书包","行李箱","笔记本电脑","校园卡","眼镜","笔记本","雨伞","水杯","其他"]
for loc, t in texts:
    if "手机" in t and "书包" in t and ("11 个" in t or "11个" in t or "11 类" in t or "11类" in t):
        missing = [c for c in CANON if c not in t]
        print(f"  {loc}: 缺={missing if missing else '无(齐全)'} | {t[:120]}")
