# -*- coding: utf-8 -*-
"""P5 论文类目名称修正前期侦察：
导出含类目名称关键词的段落（索引/样式/文本），供定位"背包/本子/玩偶"旧矛盾。
"""
from pathlib import Path
from docx import Document

DOCX = Path(r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx")
OUT = Path(r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/tools/_p5_recon_out.txt")

# 关键词：类目名称相关 + 矛盾点
KEYWORDS = [
    "背包", "书包", "本子", "笔记本", "玩偶", "娃娃",
    "类目", "类别", "物品类", "12类", "11类", "十二类", "十一类",
    "手机", "钱包", "钥匙", "行李箱", "箱包", "电脑", "校园卡", "眼镜",
    "雨伞", "水瓶", "水杯", "伞",
    "doll", "Doll", "玩偶",
]

doc = Document(str(DOCX))
paras = doc.paragraphs

lines = []
lines.append(f"# P5 侦察输出  总段数={len(paras)}  表数={len(doc.tables)}\n")

# 1) 关键词命中段落
lines.append("\n## 一、关键词命中段落\n")
seen = set()
for i, p in enumerate(paras):
    t = p.text or ""
    for kw in KEYWORDS:
        if kw in t and i not in seen:
            seen.add(i)
            lines.append(f"[{i}] <{p.style.name}> {t[:160]}")
            break

# 2) 含"类目/类别/12类"的段落附近，导出枚举上下文（前后各1段）
lines.append("\n## 二、类目枚举上下文（命中'类目/类别/12类/11类'段落 + 前后1段）\n")
enum_idx = [i for i, p in enumerate(paras) if any(k in (p.text or "") for k in ["类目", "类别", "12类", "11类", "十二类", "十一类"])]
for i in enum_idx:
    for j in range(max(0, i-1), min(len(paras), i+2)):
        p = paras[j]
        mark = ">>>" if j == i else "   "
        lines.append(f"{mark}[{j}] <{p.style.name}> {(p.text or '')[:160]}")

# 3) 表格内容（找类目枚举表）
lines.append("\n## 三、表格内容（前3张表，找类目枚举）\n")
for ti, tbl in enumerate(doc.tables[:3]):
    lines.append(f"\n### 表 {ti}\n")
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip() for c in row.cells]
        lines.append(f"  r{ri}: {cells}")

Path(OUT).write_text("\n".join(lines), encoding="utf-8")
print(f"P5 recon done -> {OUT}  (命中段落 {len(seen)} 个, 枚举段 {len(enum_idx)} 个)")
