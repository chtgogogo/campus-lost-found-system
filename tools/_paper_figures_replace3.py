# -*- coding: utf-8 -*-
"""_paper_figures_replace3.py — 将 gen3 生成的 8 张新图替换进论文 docx，并同步修改题注文字。

运行方式（建议在 tools/ 目录下）：
    python _paper_figures_replace3.py

前置条件：
    - tools/_paper_figures_new/ 下已有 8 张 PNG（由 _paper_figures_gen3.py 生成）
    - DOCX 路径指向实际论文文件

输出：
    - 原始 docx 复制为 .bak.figures（旧备份先迁移为 .bak.figures.prev）
    - 替换后 docx 直接覆盖保存
    - 控制台打印每处替换的 (图号, rId, 旧尺寸→新尺寸, 题注修改) 清单

验证：
    python _paper_figures_replace3.py --verify   （仅检查不替换）
"""
from __future__ import annotations

import os
import re
import shutil
import sys
from typing import Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# 路径配置
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
NEW_DIR = os.path.join(SCRIPT_DIR, "_paper_figures_new")

DOCX = (
    r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/"
    "曹灏天计算机学院毕业论文-2026版（7-6）.docx"
)
BAK_PATH = DOCX + ".bak.figures"

EMU_PER_CM: float = 360000.0

# ---------------------------------------------------------------------------
# 图号 → 新 PNG 映射（按题注正则匹配）
# ---------------------------------------------------------------------------
TARGETS: List[Tuple[str, str, str]] = [
    # (caption_regex, png_filename, display_label)
    (r"图\s*3\.4", "fig_03_use_case.png", "图3.4 用例图"),
    (r"图\s*3\.5", "fig_05_sequence.png", "图3.5 序列图"),
    (r"图\s*3\.7", "fig_07_state.png", "图3.7 状态图"),
    (r"图\s*3\.9", "fig_09_er.png", "图3.9 E-R图"),
    (r"图\s*3\.10", "fig_10_class.png", "图3.10 类图"),
    (r"图\s*3\.11", "fig_11_deploy.png", "图3.11 部署图"),
    (r"图\s*4\.1", "fig_41_flow.png", "图4.1 流程图"),
    (r"图\s*4\.2", "fig_42_pseudocode.png", "图4.2 伪代码"),
]

# ---------------------------------------------------------------------------
# 题注文字替换规则
# ---------------------------------------------------------------------------
TEXT_REPLACEMENTS: List[Tuple[re.Pattern, str]] = [
    # 1) "11 类" → "12 类"（全文替换，排除 "11 个...类" 以免误改 "11 个常规类"）
    (re.compile(r'(?<!个)11\s*类'), '12 类'),
    # 2) YOLOv8n → YOLOv8s（全文替换）
    (re.compile(r'YOLOv8n'), 'YOLOv8s'),
    # 3) WebSocket → 轮询（架构/部署相关描述中若有）
    (re.compile(r'WebSocket'), '轮询'),
]


def backup_docx() -> None:
    """将原始 docx 备份为 .bak.figures；若已存在则先迁移旧备份。"""
    if not os.path.exists(DOCX):
        print(f"[FATAL] docx 不存在: {DOCX}")
        sys.exit(1)
    if os.path.exists(BAK_PATH):
        prev = BAK_PATH + ".prev"
        if os.path.exists(prev):
            os.remove(prev)
        shutil.move(BAK_PATH, prev)
        print(f"[INFO] 已有备份，旧备份迁移至 {prev}")
    shutil.copy2(DOCX, BAK_PATH)
    print(f"[OK] 原始 docx 已备份至 {BAK_PATH}")


def scan_images(doc):
    """扫描文档中所有内联图片段落，返回 [(para_idx, rId, caption_text), ...]。

    caption 取图片段落后方最近一个非空段落文本（最多往后看 4 段）。
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    paras = list(body.findall(qn('w:p')))
    texts = [p.text for p in paras]

    def next_cap(i: int) -> str:
        for j in range(i + 1, min(i + 5, len(texts))):
            t = texts[j].strip()
            if t:
                return t
        return ""

    entries: List[Tuple[int, str, str]] = []
    for i, p in enumerate(paras):
        blips = list(p.iter(qn('a:blip')))
        if not blips:
            continue
        rid = blips[0].get(qn('r:embed'))
        cap = next_cap(i)
        entries.append((i, rid, cap))
    return entries


def build_plan(entries: List[Tuple[int, str, str]], doc):
    """根据题注匹配目标图号，返回 {rId: (png_path, label, orig_cx, orig_cy)}。

    同一 rId 可能被多个段落匹配到（取第一个匹配到的目标）。
    """
    plan: Dict[str, Tuple[str, str, int, int]] = {}
    for para_idx, rid, cap in entries:
        if rid in plan:
            continue  # 该 rId 已规划
        for pat, png, label in TARGETS:
            if re.search(pat, cap):
                # 提取该段落原始 extent
                cx, cy = _get_extent(para_idx, doc)
                plan[rid] = (os.path.join(NEW_DIR, png), label, cx, cy)
                break
    return plan


def _get_extent(para_idx: int, doc) -> Tuple[int, int]:
    """获取指定段落中图片的 wp:extent (cx, cy in EMU)。"""
    from docx.oxml.ns import qn
    body = doc.element.body
    paras = list(body.findall(qn('w:p')))
    if para_idx >= len(paras):
        return 0, 0
    exts = list(paras[para_idx].iter(qn('wp:extent')))
    if not exts:
        return 0, 0
    cx = int(exts[0].get('cx', '0'))
    cy = int(exts[0].get('cy', '0'))
    return cx, cy


def replace_images(plan: Dict[str, Tuple[str, str, int, int]], doc) -> List[dict]:
    """执行图片替换：更新 rel blob + 调整所有引用该 rId 的段落 extent。

    返回每条替换记录列表。
    """
    from PIL import Image
    from docx.oxml.ns import qn

    records: List[dict] = []
    body = doc.element.body
    paras = list(body.findall(qn('w:p')))

    for rid, (png_path, label, orig_cx, orig_cy) in plan.items():
        # ---- 读取新图 ----
        # 直接以原始字节写入，避免 PIL 二次编码导致与磁盘文件 md5 不一致
        # （仅作尺寸校验，不改变像素内容）。
        with Image.open(png_path) as im:
            nw, nh = im.size
        with open(png_path, "rb") as fh:
            new_bytes = fh.read()

        # ---- 替换 blob ----
        rel = doc.part.rels[rid]
        old_size = len(rel.target_part.blob)
        rel.target_part._blob = new_bytes

        # ---- 计算新 extent（保持原宽度 cx，高度按比例缩放）----
        new_cy = orig_cy
        if orig_cx > 0 and nw > 0:
            new_cy = int(orig_cx * nh / nw)

        # ---- 更新所有使用此 rId 的段落的 extent cy ----
        count = 0
        for p in paras:
            blips = list(p.iter(qn('a:blip')))
            for b in blips:
                if b.get(qn('r:embed')) == rid:
                    exts = list(p.iter(qn('wp:extent')))
                    for e in exts:
                        e.set('cy', str(new_cy))
                        # 同时确保 cx 不变（防御性）
                        if orig_cx > 0:
                            e.set('cx', str(orig_cx))
                        count += 1

        records.append({
            "label": label,
            "rid": rid,
            "old_extent": f"{orig_cx / EMU_PER_CM:.2f}cm x {orig_cy / EMU_PER_CM:.2f}cm",
            "new_img": f"{nw}x{nh}px",
            "new_extent": f"{orig_cx / EMU_PER_CM:.2f}cm x {new_cy / EMU_PER_CM:.2f}cm",
            "affected_paragraphs": count,
            "old_blob": old_size,
            "new_blob": len(new_bytes),
        })
    return records


def apply_caption_replacements(doc) -> List[Tuple[str, str]]:
    """遍历文档正文段落和表格单元格，执行 TEXT_REPLACEMENTS。返回修改清单。"""
    changes: List[Tuple[str, str]] = []

    def process_runs(container):
        for para in container.paragraphs:
            for run in para.runs:
                for pattern, repl in TEXT_REPLACEMENTS:
                    m = pattern.search(run.text)
                    if m:
                        run.text = pattern.sub(repl, run.text)
                        changes.append((m.group(), repl))

    # 正文段落
    process_runs(doc)

    # 表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_runs(cell)

    return changes


def main() -> int:
    verify_only = "--verify" in sys.argv

    # ------------------------------------------------------------------
    # 1. 导入依赖
    # ------------------------------------------------------------------
    try:
        from docx import Document
        from PIL import Image
    except ImportError as e:
        print(f"[FATAL] 缺少依赖: {e}")
        print("请使用项目 venv 运行:")
        print("  .venv\\Scripts\\python.exe _paper_figures_replace3.py")
        return 1

    # ------------------------------------------------------------------
    # 2. 检查输入
    # ------------------------------------------------------------------
    if not os.path.exists(DOCX):
        print(f"[FATAL] docx 不存在: {DOCX}")
        return 1

    missing_pngs = []
    for _, png, _ in TARGETS:
        p = os.path.join(NEW_DIR, png)
        if not os.path.exists(p):
            missing_pngs.append(p)
    if missing_pngs:
        print("[FATAL] 缺少以下新图文件：")
        for p in missing_pngs:
            print(f"  {p}")
        print("请先运行 _paper_figures_gen3.py 生成全部 8 张图。")
        return 1

    # ------------------------------------------------------------------
    # 3. 打开文档 & 扫描
    # ------------------------------------------------------------------
    doc = Document(DOCX)
    entries = scan_images(doc)
    print(f"\n== 扫描结果：共发现 {len(entries)} 个图片段落 ==")
    for pi, rid, cap in entries[:30]:
        print(f"  段[{pi:3d}] rId={rid:<8} 题注={cap!r}")
    if len(entries) > 30:
        print(f"  ... 共 {len(entries)} 个（省略后续）")

    # ------------------------------------------------------------------
    # 4. 构建替换计划
    # ------------------------------------------------------------------
    plan = build_plan(entries, doc)
    print(f"\n== 替换计划：匹配到 {len(plan)} 个唯一 rId ==")
    for rid, (path, label, cx, cy) in plan.items():
        print(f"  rId={rid:<8} → {label}  ({path})")

    if not plan:
        print("\n[WARN] 未匹配到任何目标图！请检查题注格式是否与 TARGETS 正则一致。")

    if verify_only:
        print("\n[--verify 模式，未做任何修改 --]")
        return 0

    # ------------------------------------------------------------------
    # 5. 备份
    # ------------------------------------------------------------------
    backup_docx()

    # ------------------------------------------------------------------
    # 6. 图片替换
    # ------------------------------------------------------------------
    img_records = replace_images(plan, doc)
    print(f"\n== 图片替换完成 ({len(img_records)} 条) ==")
    for rec in img_records:
        print(
            f"  [{rec['label']}] rId={rec['rid']} "
            f"extent {rec['old_extent']} → {rec['new_extent']} "
            f"(img {rec['new_img']}, {rec['affected_paragraphs']} 段落)"
        )

    # ------------------------------------------------------------------
    # 7. 题注文字替换
    # ------------------------------------------------------------------
    text_changes = apply_caption_replacements(doc)
    print(f"\n== 题注文字替换完成 ({len(text_changes)} 处) ==")
    seen = set()
    for old, new in text_changes:
        key = (old, new)
        if key not in seen:
            seen.add(key)
            cnt = sum(1 for o, n in text_changes if o == old and n == new)
            print(f"  '{old}' → '{new}'  (×{cnt})")

    # ------------------------------------------------------------------
    # 8. 保存
    # ------------------------------------------------------------------
    doc.save(DOCX)
    print(f"\n[OK] docx 已保存: {DOCX}")

    # ------------------------------------------------------------------
    # 9. 输出汇总
    # ------------------------------------------------------------------
    print("\n" + "=" * 64)
    print("  替换汇总报告")
    print("=" * 64)
    print(f"  备份文件 : {BAK_PATH}")
    print(f"  图片替换 : {len(img_records)} 个 rId")
    print(f"  文字替换 : {len(text_changes)} 处")
    print("=" * 64)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
