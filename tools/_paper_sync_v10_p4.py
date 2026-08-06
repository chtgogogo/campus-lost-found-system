# -*- coding: utf-8 -*-
"""P4 论文补丁：修复 7 类问题。

修复清单：
  ① [317] 结论段如有"稀有类 AP 偏低"表述 → 修正为实测一致
  ② [226][239] 如有"4.5节"引用 → 改为"4.4节"（目录无 4.5）
  ③ [287] 管理员段移到 [286] 本章小结之前
  ④ [192] credit_log → trust_score_log（代码真值 trust_score_log）
  ⑤ [143][211] 公式 10·other → 10·keyword（第七维代码真名 keyword）
  ⑥ table 2 COCO 通用类列头标注"(9 类)"
  ⑦ 13 处正文被错标 Heading → 还原 Normal

运行方式（托管 venv）：
  C:/Users/ASUS/.workbuddy/binaries/python/envs/default/Scripts/python.exe \
      tools/_paper_sync_v10_p4.py
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document

DOCX = Path(
    r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/"
    r"曹灏天计算机学院毕业论文-2026版（7-6）.docx"
)
BAK = DOCX.with_suffix(".docx.bak.p4")


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def replace_in_para(para, old: str, new: str) -> bool:
    """Replace *old* with *new* in a paragraph, supporting cross-run text.

    Returns True if a replacement was made.
    """
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def find_paras(doc, fragment: str):
    """Yield (index, para) for paragraphs whose text contains *fragment*."""
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            yield i, p


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------
def main():
    if not DOCX.exists():
        sys.exit(f"[错误] 论文文件不存在: {DOCX}")

    # 1. backup
    shutil.copy2(DOCX, BAK)
    print(f"[备份] {BAK}")

    # 2. load
    doc = Document(str(DOCX))
    paras = doc.paragraphs
    print(f"[加载] {len(paras)} 段, {len(doc.tables)} 表")

    fixes: list[str] = []

    # --- 修复 ④ credit_log → trust_score_log ---
    for i, p in find_paras(doc, "credit_log"):
        if replace_in_para(p, "credit_log", "trust_score_log"):
            fixes.append(f"[④] 段[{i}] credit_log → trust_score_log")

    # --- 修复 ⑤ 公式 10·other → 10·keyword ---
    for i, p in find_paras(doc, "10·other"):
        if replace_in_para(p, "10·other", "10·keyword"):
            fixes.append(f"[⑤] 段[{i}] 10·other → 10·keyword")

    # --- 诊断 ① "稀有类"相关段落 ---
    print("\n--- ① 诊断：包含'稀有类'的段落 ---")
    for i, p in find_paras(doc, "稀有类"):
        print(f"  段[{i}] ({len(p.text)}字): {p.text}")
    # 条件替换：[317] 中"仍受样本不均衡限制、AP 偏低"与实测打脸
    # 实测：campus_card 0.754、wallet 0.912、keys 0.817、glasses 0.941 全场最高
    AP_FIX_PAIRS = [
        (
            "仍受样本不均衡限制、AP 偏低，后续可扩充专属类样本与针对性数据增强进一步巩固",
            "虽样本稀少但类内形态一致、AP 反而较高（0.754~0.941），后续可进一步扩充样本巩固优势",
        ),
        ("AP 偏低", "AP 已较高"),  # 兜底
    ]
    for i, p in find_paras(doc, "稀有类"):
        changed = False
        for old, new in AP_FIX_PAIRS:
            if replace_in_para(p, old, new):
                fixes.append(f"[①] 段[{i}] 稀有类 AP 偏低 → 修正为 AP 已较高")
                changed = True
                break
        if not changed and "偏低" in p.text and "稀有类" in p.text:
            print(f"  [①警告] 段[{i}] 含'稀有类'+'偏低'但未匹配，需人工检查")

    # --- 诊断 ② "4.5"相关段落 ---
    print("\n--- ② 诊断：包含'4.5'的段落 ---")
    for i, p in enumerate(paras):
        if "4.5" in p.text:
            print(f"  段[{i}] ({len(p.text)}字): {p.text}")
    # 条件替换："4.5节"/"4.5 节" → "4.4节"/"4.4 节"（目录无 4.5，消融在 4.4）
    for i, p in find_paras(doc, "4.5"):
        if replace_in_para(p, "见 4.5 节", "见 4.4 节"):
            fixes.append(f"[②] 段[{i}] 见 4.5 节 → 见 4.4 节")
        elif replace_in_para(p, "4.5节", "4.4节"):
            fixes.append(f"[②] 段[{i}] 4.5节 → 4.4节")
        elif replace_in_para(p, "见4.5", "见4.4"):
            fixes.append(f"[②] 段[{i}] 见4.5 → 见4.4")
        elif replace_in_para(p, "详见4.5", "详见4.4"):
            fixes.append(f"[②] 段[{i}] 详见4.5 → 详见4.4")

    # --- 诊断 ⑤+ "其他类"在评分上下文 ---
    print("\n--- ⑤+ 诊断：包含'其他类'的段落（区分第七维 vs 第12类） ---")
    for i, p in enumerate(paras):
        if "其他类" in p.text and (
            "score" in p.text
            or "维度" in p.text
            or "七维" in p.text
            or "加权" in p.text
            or "keyword" in p.text
            or "other" in p.text.lower()
        ):
            print(f"  段[{i}] ({len(p.text)}字): {p.text}")

    # --- 修复 ⑥ table 2 COCO 通用类列头 ---
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if txt == "COCO 通用类" and "9" not in txt:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = "COCO 通用类（9 类）"
                        fixes.append(
                            f"[⑥] 表{ti} r{ri}c{ci} "
                            "COCO 通用类 → COCO 通用类（9 类）"
                        )

    # --- 修复 ⑦ 13 处标题样式 ---
    HEADING_FIX_PATTERNS = [
        "匹配模块（MatchService）是算法核心",
        "候选集先经类别主键过滤",
        "发布模块（PublishService）",
        "交接与审计模块（HandoverService",
        "本章明确了以 YOLOv8s + YOLO-World",
        "系统实现采用既定技术栈",
        "核心模块实现思路如下",
        "本章在既定技术栈下完成核心模块编码",
        "测试环境：Windows 平台",
        "测试设计遵循",
        "关键模块测试记录如表 5-1",
        "测试运行结果：全量 334",
        "本章通过单元与端到端测试对系统验证",
    ]
    for i, p in enumerate(paras):
        if p.style.name.startswith("Heading"):
            for pattern in HEADING_FIX_PATTERNS:
                if pattern in p.text:
                    old_style = p.style.name
                    p.style = doc.styles["Normal"]
                    fixes.append(
                        f"[⑦] 段[{i}] {old_style} → Normal: {p.text[:50]}"
                    )
                    break

    # --- 修复 ③ 管理员段移到本章小结之前 ---
    admin_idx = -1
    summary_idx = -1
    for i, p in enumerate(paras):
        if "管理员能力在 v7" in p.text or "管理员能力在v7" in p.text:
            admin_idx = i
        if "本章在既定技术栈下完成核心模块编码" in p.text:
            summary_idx = i

    if admin_idx >= 0 and summary_idx >= 0 and admin_idx > summary_idx:
        admin_elem = paras[admin_idx]._element
        summary_elem = paras[summary_idx]._element
        summary_elem.addprevious(admin_elem)
        fixes.append(
            f"[③] 管理员段[{admin_idx}] 移到 小结段[{summary_idx}] 之前"
        )
    else:
        print(
            f"[③诊断] admin_idx={admin_idx}, summary_idx={summary_idx}"
        )
        if admin_idx >= 0:
            print(f"  管理员段: {paras[admin_idx].text[:80]}")
        if summary_idx >= 0:
            print(f"  小结段: {paras[summary_idx].text[:80]}")
        if admin_idx >= 0 and summary_idx >= 0 and admin_idx < summary_idx:
            print("  [③信息] 管理员段已在小结段之前，无需移动")

    # 3. save
    doc.save(str(DOCX))

    # 4. print fix list
    print(f"\n=== P4 修复清单 ({len(fixes)} 项) ===")
    for f in fixes:
        print(f"  {f}")

    # 5. self-check
    print("\n=== 自检 ===")
    doc2 = Document(str(DOCX))
    paras2 = doc2.paragraphs

    # check ④
    credit_count = sum(1 for p in paras2 if "credit_log" in p.text)
    print(f"  credit_log 残留: {credit_count} (应为 0)")

    # check ⑤
    other_count = sum(1 for p in paras2 if "10·other" in p.text)
    print(f"  10·other 残留: {other_count} (应为 0)")

    # check ⑦
    heading_fix_count = 0
    for p in paras2:
        if p.style.name.startswith("Heading"):
            for pattern in HEADING_FIX_PATTERNS:
                if pattern in p.text:
                    heading_fix_count += 1
                    break
    print(f"  标题样式残留: {heading_fix_count} (应为 0)")

    # check ③
    admin_pos = -1
    summary_pos = -1
    for i, p in enumerate(paras2):
        if "管理员能力在 v7" in p.text or "管理员能力在v7" in p.text:
            admin_pos = i
        if "本章在既定技术栈下完成核心模块编码" in p.text:
            summary_pos = i
    if admin_pos >= 0 and summary_pos >= 0:
        ok = admin_pos < summary_pos
        print(
            f"  管理员段[{admin_pos}] < 小结段[{summary_pos}]: {ok} "
            f"(应为 True)"
        )
    else:
        print(
            f"  管理员段/小结段定位: admin={admin_pos}, summary={summary_pos}"
        )

    # check ⑥
    for ti, table in enumerate(doc2.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if "COCO 通用类" in cell.text and "9" in cell.text:
                    print(f"  表{ti} COCO 通用类（9 类）: OK")

    print("\n[P4] 完成")


if __name__ == "__main__":
    main()
