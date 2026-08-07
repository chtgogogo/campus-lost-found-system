import os
from docx import Document
from docx.oxml.ns import qn

NEW = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_paper_figures_new")
DOCX = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
EMU_PER_CM = 360000.0

TARGETS = {"rId20":"fig_03_use_case.png","rId21":"fig_05_sequence.png","rId22":"fig_07_state.png",
 "rId23":"fig_11_deploy.png","rId24":"fig_10_class.png","rId27":"fig_05_sequence.png",
 "rId30":"fig_09_er.png","rId31":"fig_10_class.png","rId32":"fig_11_deploy.png",
 "rId28":"fig_41_flow.png","rId35":"fig_42_pseudocode.png"}

# A4 text area width ~ 16-17cm (margins ~2cm each side on 21cm page)
PAGE_TEXT_W_CM = 16.5

doc = Document(DOCX)
body = doc.element.body
paras = list(body.findall(qn('w:p')))
print("== 各目标图版心宽度 / 比例校验 ==")
all_safe = True
for rid, png in TARGETS.items():
    # find all extents referencing rid
    cxs, cys = [], []
    for p in paras:
        for b in p.iter(qn('a:blip')):
            if b.get(qn('r:embed')) == rid:
                for e in p.iter(qn('wp:extent')):
                    try:
                        cxs.append(int(e.get('cx','0'))); cys.append(int(e.get('cy','0')))
                    except: pass
    if not cxs:
        print(f"  {rid:8} 未找到 extent"); continue
    cx = cxs[0]; cy = cys[0]
    wcm = cx / EMU_PER_CM
    # aspect ratio check
    from PIL import Image
    im = Image.open(os.path.join(NEW, png)); iw, ih = im.size
    img_ratio = iw/ih
    ext_ratio = cx/cy
    ratio_ok = abs(img_ratio - ext_ratio) < 0.02
    safe = wcm <= PAGE_TEXT_W_CM
    all_safe = all_safe and safe and ratio_ok
    print(f"  {rid:8} 图宽={wcm:6.2f}cm 比例img={img_ratio:.3f}/ext={ext_ratio:.3f} "
          f"{'比例OK' if ratio_ok else '*** 比例失真 ***'} {'宽度OK' if safe else '*** 超宽 ***'}")
print(f"全部版心安全且比例正确: {all_safe}")

# 全文关键字复检
print("\n== 全文关键字复检 ==")
full = ""
for p in doc.paragraphs: full += p.text + "\n"
for t in doc.tables:
    for r in t.rows:
        for c in r.cells: full += c.text + "\n"
for kw in ["YOLOv8n","YOLOv8s","WebSocket","轮询","11 类","12 类","11类","12类"]:
    print(f"  {kw!r:12}: {full.count(kw)}")
