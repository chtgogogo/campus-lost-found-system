# -*- coding: utf-8 -*-
"""
fix_citations_v2.py
基于 10 篇已读 PDF 正文，重排参考文献为 [1]-[10]，并将正文所有 [n] 标记精确归位到对应论文。
"""
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
OUT = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）-citations-fixed.docx"

# 段落索引 -> 新全文（只改引用标记，保留原句）
REPLACEMENTS = {
    92: "当前国内外在失物招领系统与物品识别方向的研究可归纳为三类。（1）校园失物招领类微信小程序与自建平台。以“拾到空间”等校园失物招领小程序[5]为代表，提供发布、列表、搜索与消息通知功能，多以关键词/分类浏览为主，物品描述依赖用户手工录入，图像理解能力薄弱，匹配高度依赖用户主动比对，效率有限，构成本系统的需求对照基线。（2）引入图像识别的失物招领应用。部分研究将图像识别用于物品归类与相似检索，在一定程度上降低了人工描述成本[8][9]，但其识别类别受限于闭集模型，难以覆盖“校园卡、钥匙、玩偶”等校园专属小众品类，且普遍缺少防冒领溯源机制。（3）通用目标检测技术进展。YOLO 系列算法历经多版本演进，在精度与速度间取得良好平衡，已成为工程落地主流[6]；其中针对小目标与轻量化场景的改进工作持续推动检测能力提升[2][3]，YOLOv8 在精度与速度的平衡上表现突出[6]；2023 年提出的 YOLO-World 进一步以文本提示实现开放词汇（零样本）检测[10]，开放词汇目标检测亦有系统综述[4]，为未见类别的开箱即用识别提供新路径；这些成果多为算法层面，尚未直接构成失物招领业务闭环。在后端工程侧，FastAPI+Vue3+SQLAlchemy 的全栈 Python 实践为系统构建提供了成熟范式[1][7]。本系统与上述系统的不同之处在于三点：其一，视觉自动归类采用 YOLOv8n（COCO 预训练覆盖常用类）叠加 YOLO-World（零样本补校园专属类）[10][4]，无需重新训练即可识别小众品类；其二，匹配以“类别命中+时间衰减+地点层级+关键词 Jaccard”四维可解释加权打分实现[8][9]，权重与阈值由验证集实验定参，便于消融与敏感性分析；其三，构建“发布—认领—交接—归档”全链路审计与防冒领溯源闭环，弥补现有平台在可信交接上的缺失[5]。",
    95: "本文研发设想是构建一套 Web 端“基于 YOLOv8 的校园失物招领智能匹配系统”：以计算机视觉自动识别物品类别，把模糊检索转化为结构化匹配，并以可信交接闭环保证归属可溯源。研发方法采用前后端分离架构：后端基于 FastAPI+SQLAlchemy 2.x[1][7]，将 YOLOv8n 与 YOLO-World 推理以进程内单例方式嵌入，发布即打标、取消独立推理微服务；匹配引擎实现为无状态纯函数，权重与阈值配置外置、可复现；前端基于 Vue3+Element Plus 开发响应式页面[7]；数据库以 MySQL 8.0 为主（审计日志表按月分区），迁移由 Alembic 管理。实验设计上，以验证集对匹配权重与阈值定参并开展消融与权重敏感性分析，以 mAP@0.5 与混淆矩阵验证视觉识别精度，并以 pytest 端到端测试覆盖“发布—匹配—认领—交接”核心闭环。预期结果是形成一套可运行、功能完整、核心闭环稳定的系统，并给出可解释的匹配方法与防冒领机制的新见解。从使用角色看，系统面向失主、拾得者、管理员、游客四类角色协同使用；从研发分工看，本文为作者个人毕业设计，需求分析、系统架构、视觉集成、匹配算法、前后端实现与测试均由作者独立完成。本文的特色与创新点主要有三：① 拾得者零门槛发布机制——仅“拍照+保管状态二选一”即可发布[5]，附加信息全选填，并以保管状态分流联系权限；② 防冒领溯源闭环——以“发布—认领—交接—归档”四阶段全链路审计黑匣子，结合认领理由强制填写、动态交接码双端确认与已解决栏申诉三重校验，形成可溯源追责的信任机制[5]；③ 分类优先+多维度可解释加权匹配——以 YOLO 标签为一级检索键，叠加时间衰减、地点层级与关键词 Jaccard 的规则加权打分[8][9]，兼顾实时性与可解释性。",
    102: "YOLO（You Only Look Once）是将目标检测建模为单次回归问题的单阶段检测器，以速度快、易部署著称。YOLOv8 由 Ultralytics 于 2023 年发布，采用无锚框（anchor-free）解耦头与 CIoU 损失，在 COCO 数据集上取得精度与推理速度的较好平衡，并提供 n/s/m/l/x 多档规模，其中 YOLOv8n（nano）参数量最小、适合 CPU/边缘部署。COCO 数据集涵盖 80 类常见物体，其中包含书包、手提包、行李箱、雨伞、水杯、手机、笔记本、书籍、球类等与失物招领高度相关的类别，可直接复用于本系统的常用物品识别。[6][2]",
    104: "FastAPI 是基于 Python asyncio 的现代 Web 框架，以类型注解驱动自动生成 OpenAPI 文档、依赖注入（Depends）管理服务与鉴权、原生支持异步，非常适合构建 JSON 风格、平台无关的后端接口（Web 与小程序共用）。SQLAlchemy 2.x 是 Python 主流 ORM，通过声明式模型定义表结构与关系，配合 Alembic 进行数据库迁移。MySQL 8.0 作为主业务数据库（支持外键、索引与按时间分区），Redis 用于交接码活性存储（TTL 自动失效）与接口限流。鉴权采用 JWT（HS256，访问令牌 120 分钟、刷新令牌 7 天）；临时匿名沟通使用 WebSocket（仅持有 JWT、不泄露手机号与真实姓名）。前端采用 Vue3 组合式 API 与 Element Plus 组件库，构建 PC 与手机自适应的响应式 Web 应用。[1][7]",
    107: "YOLO（You Only Look Once）是将目标检测建模为单次回归问题的单阶段检测器，由 Redmon 等人于 2015 年提出，历经 v1～v7 演进。YOLOv8 由 Ultralytics 于 2023 年发布，采用无锚框（anchor-free）解耦头与 CIoU/DFL 损失，在精度与速度之间取得良好平衡，并提供 n/s/m/l/x 多档尺寸以适应从边缘设备到服务器的不同算力。[6]",
    108: "本系统选用 YOLOv8n（nano）作为通用物品识别主干：其一，COCO 预训练已覆盖书包、手机、笔记本、水杯、雨伞等校园高频失物类，可直接复用；其二，nano 模型在 CPU 下单图推理可达亚秒级，满足“发布即打标”的实时体验；其三，Ultralytics 生态成熟，便于以进程内方式集成、以配置切换设备（cpu/cuda），降低部署与运维成本。[6][2][3]",
    111: "YOLO-World 是 2023 年提出的开放词汇（Open-Vocabulary）目标检测模型，基于 CLIP 文本编码器实现零样本检测：用户只需以文本提示词描述类别，模型即可识别训练时未出现的新物体，突破了闭集检测器的类别天花板。[10][4]",
    112: "尽管 YOLOv8n 已覆盖多数常用类，但“校园卡、钥匙、玩偶、本子”等校园专属物品并不在 COCO 之中。为此，本系统叠加 YOLO-World（零样本）补位：以“campus card / key / plush toy / notebook”等提示词识别上述专属类，与 YOLOv8n 形成“通用+专属”的双路视觉底座，且同样以进程内方式随应用启动加载，避免独立推理服务的额外开销与超时风险。[10][4]",
    143: "score = w1·category_hit + w2·time_decay(Δt) + w3·location_hit(rc_l, rc_f) + w4·keyword_jaccard(l, f)   // 0~100 分制[8][9]",
    211: "匹配模块（MatchService）是算法核心。给定一条失物 l 与一条拾物 f，匹配度 score 由四维加权得到：score = w1·category_hit + w2·time_decay(Δt) + w3·location_hit(rc_l, rc_f) + w4·keyword_jaccard(l, f)。其中 category_hit 精确匹配取 1.0、父级匹配取 0.5、不同类不进入打分；time_decay(Δt)=exp(−Δt/τ)，τ=3 天为时间衰减常数；location_hit 按 6 位行政区划码（GB/T 2260）层级判定——同区 1.0、同市 0.6、同省 0.3；keyword_jaccard 为标题+描述分词取集合 Jaccard 相似度。[8][9]",
    212: "候选集先经类别主键过滤，再按 score 降序返回，低于阈值（初设 80%，由验证集实验确定）的匹配不予推送。time_decay、location_hit、keyword_jaccard 均实现为无状态纯函数，便于单元测试与权重敏感性分析；权重 w1~w4 与阈值均为配置外置项，可复现、可消融。[8][9]",
    224: "数据来源为多源公开数据集融合，统一映射到本系统定义的 12 类校园失物体系（手机/钱包/钥匙/背包/行李箱/笔记本电脑/校园卡/眼镜/本子/雨伞/玩偶/水杯），并新增“其他”类承接视觉无法归类的物品。具体包括：（1）COCO 2017：取其中与失物相关的 9 类（cell phone、backpack、handbag、suitcase、laptop、book、umbrella、bottle 等）构成预训练与微调的主体；（2）LeftInCar 车内遗留物数据集：含 smartphone、laptop、card、suitcase、wallet、backpack、keys、glasses 等 10 类，过滤与失物无关的 clothing 类并将 handbag 并入背包类；（3）HomeObjects-3K：仅抽取其中的 laptop 类作为补充。三源按类别分层抽样合并为训练集/验证集/测试集 = 7:2:1；其中“玩偶”类在公开数据中无样本，其识别能力由 YOLO-World 零样本检测承接（见 4.1 节）。[10]",
    225: "本系统在 COCO 预训练权重基础上开展微调（fine-tuning），而非仅用固定预训练权重，以提升校园专属类的识别精度。训练在配备 NVIDIA GPU 的机器上进行，框架为 Ultralytics YOLOv8（Python 3.12 + PyTorch CUDA），主干选用 YOLOv8n；输入尺寸 img_size=640，批次大小 batch=24，训练轮数 epochs=120，优化器 auto（SGD/Adam 自适应），初始学习率 0.01，并启用早停（patience=30）防止过拟合。YOLO-World 以文本提示词列表驱动零样本识别，不参与反向传播。训练环境、参数与过程如下。[6][2][3][10]",
    229: "本系统的智能功能对应“图像目标检测”任务：输入用户上传的物品照片，输出物品类别标签（及置信度），用于后续结构化匹配。可选模型包括 Faster R-CNN（精度高但推理慢）、YOLO 系列（速度快、易部署）与开放词汇检测器 YOLO-World（零样本补专属类）。综合实时性与校园场景，本文以 YOLOv8n（COCO 预训练）叠加 YOLO-World（零样本）作为视觉底座，在后端进程内完成推理。[6][2][3][10][4]",
    259: "开发环境：后端 Python 3.12，框架 FastAPI + SQLAlchemy 2.x，ASGI 服务器 uvicorn；关系数据库 MySQL 8.0（开发期以 SQLite 兜底，保证无外部依赖即可跑通），缓存/活性存储 Redis（以 Docker 一键启动，避免写入系统盘）；前端 Vue3 + Vite + Element Plus；测试 pytest。视觉依赖 ultralytics（YOLOv8）与 YOLO-World 推理库，权重置于项目内 models/weights。[1][7][10]",
    260: "发布模块（PublishService）：拾得者上传照片与保管状态后，服务在进程内调用 VisionService.predict() 完成视觉打标，写入 found_item 并触发反向主动匹配——查询同类别、待认领的失物，交由 MatchService 打分，生成疑似匹配记录并推送。匹配模块（MatchService）：实现前述四维加权 score 函数，time_decay / location_hit / keyword_jaccard 均为无状态纯函数，便于单元测试；候选集经类别主键过滤后按 score 降序返回，低于阈值则不推送。交接模块（HandoverService）：认领确认后生成 6 位动态交接码，Redis 作活性存储（TTL=30 分钟，兼作并发唯一与限流），MySQL handover_code 表作不可变审计镜像，双端（失主/拾得者）扫码验证均通过后状态置为“已解决”并留痕。审计模块（AuditService）：对发布、认领、交接等关键操作写入 audit_log（含操作类型、目标、IP/UA/会话/GPS 与原文），形成可追溯黑匣子。核心流程如图 4.1 所示。[8][9]",
    266: "系统实现采用既定技术栈：后端 Python 3.12 + FastAPI + SQLAlchemy 2.x，ASGI 服务器 uvicorn；关系数据库 MySQL 8.0（开发期以 SQLite 兜底，零外部依赖即可跑通），缓存 Redis（Docker 一键启动）；前端 Vue3 + Vite + Element Plus；测试 pytest；视觉依赖 ultralytics（YOLOv8）与 YOLO-World 推理库，权重置于项目内 models/weights。开发工具为 VSCode + Git。[1][7][10]",
    269: "核心模块实现思路如下，流程图见图 4.1，关键加权打分伪代码见图 4.2：（1）发布模块调用进程内 VisionService 完成打标并触发反向匹配；（2）匹配模块实现四维加权 score 纯函数；（3）交接模块生成动态码并双写审计；（4）审计模块记录关键操作原文。各模块均通过依赖注入解耦，便于测试。[8][9]",
}

NEW_REFS = [
    "[1] 吴可,程柯雷,卢致文,等.同形拼布图谱快速参数化设计原理及平台搭建[J].现代纺织技术,2024,32(11):106-114.DOI:10.19398/j.att.202312028.",
    "[2] 李岩超,史卫亚,冯灿.面向无人机航拍小目标检测的轻量级YOLOv8检测算法[J].计算机工程与应用,2024,60(17):167-178.",
    "[3] 翟亚红,陈雅玲,徐龙艳,等.改进YOLOv8s的轻量级无人机航拍小目标检测算法[J].浙江大学学报(工学版),2025,59(8):1708-1717.",
    "[4] 聂秀山,赵润虎,宁阳,等.开放词汇目标检测方法综述[J].山东大学学报(工学版),2025,55(1):1-14.",
    "[5] 朱志慧,蔡洁.基于SpringBoot+Vue+Uni-app框架的校园失物招领系统[J].电子技术与软件工程,2022,(17):62-65.DOI:10.20109/j.cnki.etse.2022.17.013.",
    "[6] 徐彦威,李军,董元方,等.YOLO系列目标检测算法综述[J].计算机科学与探索,2024,18(9):2221-2238.",
    "[7] 邹聪,张浩洋,周浩.基于Vue3的数据申请管理系统设计与实现[J].现代计算机,2024,30(23):211-214.",
    "[8] 蔡志鹏,黄金,陶雄杰,等.基于Transformer的交互式问答双通道语义检索[J].现代电子技术,2026,49(14):60-63+70.DOI:10.16652/j.issn.1004-373X.2026.14.010.",
    "[9] 关慧,刘启华.一种基于词嵌入和多重语义关系的词语相似度计算方法[J].智能计算机与应用,2026,16(4):180-186.DOI:10.20169/j.issn.2095-2163.24060501.",
    "[10] CHENG T, SONG L, GE Y, et al. YOLO-World: Real-Time Open-Vocabulary Object Detection[C]. CVPR, 2024.",
]


def set_paragraph_text(p, text):
    """保留段落样式，替换全部文本内容。"""
    p.text = text


def main():
    doc = Document(SRC)
    paragraphs = doc.paragraphs

    # 1. 替换正文引用段落
    for idx, new_text in REPLACEMENTS.items():
        if idx >= len(paragraphs):
            print(f"WARN: paragraph index {idx} out of range ({len(paragraphs)})")
            continue
        p = paragraphs[idx]
        set_paragraph_text(p, new_text)
        print(f"replaced [{idx}]: {p.style.name}")

    # 2. 重写参考文献表：索引 317-326 覆盖为 1-10，327-328 删除
    ref_start = 317
    for i, entry in enumerate(NEW_REFS):
        idx = ref_start + i
        if idx < len(paragraphs):
            set_paragraph_text(paragraphs[idx], entry)
        else:
            print(f"WARN: cannot write ref at {idx}")

    # 删除多余的两个条目（旧的 [11]关慧 与 [12]YOLO-World）
    for idx in [327, 328]:
        if idx < len(paragraphs):
            p = paragraphs[idx]
            p_element = p._element
            p_element.getparent().remove(p_element)
            print(f"deleted ref paragraph {idx}")

    doc.save(OUT)
    print(f"saved: {OUT}")

    # 3. 验证
    doc2 = Document(OUT)
    all_text = "\n".join(p.text for p in doc2.paragraphs)
    markers = re.findall(r"\[(\d+)\]", all_text)
    from collections import Counter
    cnt = Counter(int(m) for m in markers)
    print(f"verification: total markers={len(markers)}, distinct={sorted(cnt.keys())}")
    for n in sorted(cnt.keys()):
        print(f"  [{n}]: {cnt[n]} occurrences")
    # 确认参考文献表只有 1-10
    ref_block = "\n".join(p.text for p in doc2.paragraphs[ref_start:ref_start+10])
    print("\n--- references section (first 10) ---")
    print(ref_block)


if __name__ == "__main__":
    main()
