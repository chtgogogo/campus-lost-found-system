# -*- coding: utf-8 -*-
"""P4 论文补丁终审 QA 脚本。

用 python-docx 重新导出修复后的 docx（336 段 + 6 表），与 P4 修复清单逐项核对。
仅审查、不修改任何文件。
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

DOCX = Path(
    r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/"
    r"曹灏天计算机学院毕业论文-2026版（7-6）.docx"
)

PY = sys.executable


def main():
    if not DOCX.exists():
        sys.exit(f"[错误] 论文文件不存在: {DOCX}")

    doc = Document(str(DOCX))
    paras = doc.paragraphs
    tables = doc.tables
    print(f"[加载] {len(paras)} 段, {len(tables)} 表\n")

    results = []  # (id, desc, PASS/FAIL, detail)

    def add(idx, desc, ok, detail=""):
        results.append((idx, desc, "PASS" if ok else "FAIL", detail))

    # ============================================================
    # ① 段[317] 结论：稀有类 AP 偏低 → AP 已较高/反而较高
    # ============================================================
    p317 = paras[317].text if len(paras) > 317 else ""
    has_ap_low = "AP 偏低" in p317
    has_ap_high = ("AP 反而较高" in p317) or ("AP 已较高" in p317)
    has_range = ("0.754" in p317) and ("0.941" in p317)
    add("①-1", "段[317] 无'AP 偏低'", not has_ap_low,
        f"含'AP 偏低'={has_ap_low}")
    add("①-2", "段[317] 含 AP 较高表述", has_ap_high,
        f"AP 反而较高={'AP 反而较高' in p317}, AP 已较高={'AP 已较高' in p317}")
    add("①-3", "段[317] 含实测区间 0.754~0.941", has_range,
        f"0.754={'0.754' in p317}, 0.941={'0.941' in p317}")

    # ============================================================
    # ② 段[226][239] 见 4.5 节 → 见 4.4 节
    # ============================================================
    for idx in (226, 239):
        pt = paras[idx].text if len(paras) > idx else ""
        has_45 = ("4.5 节" in pt) or ("4.5节" in pt)
        has_44 = ("4.4 节" in pt) or ("4.4节" in pt)
        add(f"②-{idx}", f"段[{idx}] 无'4.5 节'且含'4.4 节'",
            (not has_45) and has_44,
            f"4.5节={has_45}, 4.4节={has_44}")

    # ============================================================
    # ③ 管理员段在本章小结段之前
    # ============================================================
    admin_idx = -1
    summary_idx = -1
    for i, p in enumerate(paras):
        if "管理员能力在 v7" in p.text or "管理员能力在v7" in p.text:
            admin_idx = i
        if "本章在既定技术栈下完成核心模块编码" in p.text:
            summary_idx = i
    ok_order = (admin_idx >= 0 and summary_idx >= 0 and admin_idx < summary_idx)
    add("③-1", "管理员段定位(<小结段)", ok_order,
        f"admin_idx={admin_idx}, summary_idx={summary_idx}")

    # ============================================================
    # ④ 段[192] credit_log → trust_score_log
    # ============================================================
    p192 = paras[192].text if len(paras) > 192 else ""
    has_credit = "credit_log" in p192
    has_trust = "trust_score_log" in p192
    add("④-1", "段[192] 无'credit_log'", not has_credit,
        f"credit_log={has_credit}")
    add("④-2", "段[192] 含'trust_score_log'", has_trust,
        f"trust_score_log={has_trust}")

    # ============================================================
    # ⑤ 段[143][211] 公式 10·other → 10·keyword
    # ============================================================
    for idx in (143, 211):
        pt = paras[idx].text if len(paras) > idx else ""
        has_other = "10·other" in pt
        has_kw = "10·keyword" in pt
        add(f"⑤-{idx}", f"段[{idx}] 无'10·other'且含'10·keyword'",
            (not has_other) and has_kw,
            f"10·other={has_other}, 10·keyword={has_kw}")

    # ============================================================
    # ⑥ 表 2 r0c1 COCO 通用类 → COCO 通用类（9 类）
    # ============================================================
    # table 2 is the mAP results table (4 rows x 5 cols)
    t2 = tables[2]
    r0c1 = t2.rows[0].cells[1].text.strip()
    add("⑥-1", "表2 r0c1 = 'COCO 通用类（9 类）'",
        r0c1 == "COCO 通用类（9 类）", f"实际='{r0c1}'")

    # ============================================================
    # ⑦ 13 处被错标的段落 style → Normal
    # ============================================================
    HEADING_FIX_PATTERNS = [
        "匹配模块（MatchService）是算法核心",
        "候选集先经类别主键过滤",
        "发布模块（PublishService）",
        "交接与审计模块（HandoverService",
        "本章明确了以 YOLOv8s + YOLO-World",
        "系统实现采用既定技术栈",
        "核心模块实现思路如下",
        "本章在既定技术栈下完成核心模块编码",
        "测试环境：Windows 平台",
        "测试设计遵循",
        "关键模块测试记录如表 5-1",
        "测试运行结果：全量 334",
        "本章通过单元与端到端测试对系统验证",
    ]
    heading_residue = 0
    normal_count = 0
    for i, p in enumerate(paras):
        for pattern in HEADING_FIX_PATTERNS:
            if pattern in p.text:
                if p.style.name.startswith("Heading"):
                    heading_residue += 1
                    add(f"⑦-p{i}", f"段[{i}] Normal(残留Heading)",
                        False, f"style={p.style.name}")
                else:
                    normal_count += 1
                    add(f"⑦-p{i}", f"段[{i}] Normal", True,
                        f"style={p.style.name}")
                break
    # 注：部分模式是通用子串，会同时命中"待修复段"与"原本就是 Normal 的章节导引段"
    # (如"测试环境：Windows 平台"同时匹配 p291 导引段与 p296 待修复段)。
    # 真正的验收标准 = 残留Heading==0 且 13 个原始 Heading 段均已 Normal。
    EXPECTED_FIX_INDICES = {
        211, 212, 216, 217, 256, 268, 271, 287,
        296, 299, 302, 308, 313,
    }
    fixed_all_normal = all(
        paras[i].style.name == "Normal" for i in EXPECTED_FIX_INDICES
        if i < len(paras)
    )
    add("⑦-summary", "13 处原始Heading段均Normal + 0残留",
        heading_residue == 0 and fixed_all_normal,
        f"Normal匹配={normal_count}(含3导引段), 残留Heading={heading_residue}, "
        f"13段全Normal={fixed_all_normal}")

    # ============================================================
    # 全局残留检查（防漏网）
    # ============================================================
    credit_global = sum(1 for p in paras if "credit_log" in p.text)
    other_global = sum(1 for p in paras if "10·other" in p.text)
    ap_low_global = sum(1 for p in paras if "AP 偏低" in p.text)
    sec45_global = sum(1 for p in paras if ("4.5 节" in p.text) or ("4.5节" in p.text))

    print("=== 全局残留 ===")
    print(f"  credit_log 残留: {credit_global} (应为 0)")
    print(f"  10·other 残留: {other_global} (应为 0)")
    print(f"  AP 偏低 残留: {ap_low_global} (应为 0)")
    print(f"  4.5节 残留: {sec45_global} (应为 0)")

    add("G-credit", "全局 credit_log 残留=0", credit_global == 0,
        f"count={credit_global}")
    add("G-other", "全局 10·other 残留=0", other_global == 0,
        f"count={other_global}")
    add("G-aplow", "全局 AP 偏低 残留=0", ap_low_global == 0,
        f"count={ap_low_global}")
    add("G-sec45", "全局 4.5节 残留=0", sec45_global == 0,
        f"count={sec45_global}")

    # ============================================================
    # 一致性检查：论文关键数据 vs 代码/评测
    # ============================================================
    print("\n=== 一致性检查 ===")
    # mAP@0.5=0.654 测试集 (表2 r1c3 整体)
    t2_r1c3 = t2.rows[1].cells[3].text.strip()
    add("C-map654", "表2 整体 mAP@0.5=0.654", t2_r1c3 == "0.654",
        f"实际='{t2_r1c3}'")

    # mAP@0.5=0.710 验证集 (段[317] 含 0.710)
    p317_has_710 = "0.710" in p317
    add("C-map710", "段[317] 验证集 mAP@0.5=0.710", p317_has_710,
        f"含0.710={p317_has_710}")

    # 334 测试用例
    p294 = paras[294].text if len(paras) > 294 else ""
    p308 = paras[308].text if len(paras) > 308 else ""
    has_334 = ("334" in p294) or ("334" in p308)
    add("C-334", "测试用例 334 提及", has_334,
        f"294含334={'334' in p294}, 308含334={'334' in p308}")

    # 12 类 nc (段[224] 含 nc=12)
    p224 = paras[224].text if len(paras) > 224 else ""
    has_nc12 = ("nc=12" in p224) or ("nc = 12" in p224)
    add("C-nc12", "段[224] nc=12", has_nc12,
        f"nc=12={has_nc12}")

    # trust_score_log 表名 (表0 r9)
    t0_r9 = tables[0].rows[9].cells[0].text.strip()
    add("C-tslog", "表0 r9 trust_score_log", t0_r9 == "trust_score_log",
        f"实际='{t0_r9}'")

    # 表2 校园专属类 mAP 0.848 (与结论 0.754~0.941 一致性)
    t2_r1c2 = t2.rows[1].cells[2].text.strip()
    add("C-campus", "表2 校园专属类 mAP=0.848", t2_r1c2 == "0.848",
        f"实际='{t2_r1c2}'")

    # ============================================================
    # 输出报告
    # ============================================================
    print("\n" + "=" * 70)
    print("P4 终审逐项核对结果")
    print("=" * 70)
    n_pass = 0
    n_fail = 0
    for idx, desc, status, detail in results:
        mark = "✅" if status == "PASS" else "❌"
        print(f"  {mark} [{status}] {idx}: {desc}  | {detail}")
        if status == "PASS":
            n_pass += 1
        else:
            n_fail += 1
    print("-" * 70)
    print(f"  合计: {len(results)} 项 | PASS: {n_pass} | FAIL: {n_fail}")
    print("=" * 70)


if __name__ == "__main__":
    main()
