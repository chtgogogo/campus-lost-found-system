"""Fill the real v4 validation mAP into the thesis results table (table index 2).

Grouping (matches the thesis data table [1]):
  COCO 通用类 (8): phone, backpack, suitcase, laptop, notebook, umbrella, bottle, glasses
  校园专属类 (3): wallet, keys, campus_card
"""
import json
import shutil

import docx

THESIS = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
METRICS = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/v4_val_metrics.json"


def set_cell(cell, val):
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = val
    else:
        cell.text = val


def main():
    d = json.load(open(METRICS, encoding="utf-8"))
    m50 = d["per_class_map50"]
    m5095 = d["per_class_map50_95"]
    coco = [0, 3, 4, 5, 8, 9, 10, 7]
    camp = [1, 2, 6]
    coco50 = sum(m50[i] for i in coco) / 8
    coco5095 = sum(m5095[i] for i in coco) / 8
    camp50 = sum(m50[i] for i in camp) / 3
    camp5095 = sum(m5095[i] for i in camp) / 3

    # backup
    bak = THESIS + ".bak.20260724"
    shutil.copy2(THESIS, bak)

    doc = docx.Document(THESIS)
    t = doc.tables[2]
    # row0 = header, row1 = mAP@0.5, row2 = mAP@0.5:0.95, row3 = confusion matrix
    set_cell(t.rows[1].cells[1], "%.3f" % coco50)
    set_cell(t.rows[1].cells[2], "%.3f" % camp50)
    set_cell(t.rows[1].cells[3], "%.3f" % d["map50"])
    set_cell(t.rows[2].cells[1], "%.3f" % coco5095)
    set_cell(t.rows[2].cells[2], "%.3f" % camp5095)
    set_cell(t.rows[2].cells[3], "%.3f" % d["map50_95"])
    doc.save(THESIS)

    print("Filled thesis table[2] with v4 mAP. Backup ->", bak)
    print("mAP@0.5      COCO=%.3f  campus=%.3f  overall=%.3f" % (coco50, camp50, d["map50"]))
    print("mAP@0.5:0.95 COCO=%.3f  campus=%.3f  overall=%.3f" % (coco5095, camp5095, d["map50_95"]))


if __name__ == "__main__":
    main()
