import docx, shutil
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r'D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）-12cls.docx'
ORIG = r'D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx'
BAK = r'D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx.bak.20260725'

# 1) backup original (11-class version) before overwrite
shutil.copy2(ORIG, BAK)
print('backup original ->', BAK, '(', __import__('os').path.getsize(BAK), 'bytes )')

# 2) load the 12-class filled copy
d = docx.Document(SRC)
t = d.tables[2]

# 3) insert an honest note paragraph right after table[2]
tbl = t._tbl
new_p = OxmlElement('w:p')
# copy default paragraph properties so it looks normal
tbl.addnext(new_p)
para = Paragraph(new_p, d)
para.text = ('注：第 12 类“其他类（other）”为无法归入前述 11 个具体失物类别的兜底类别，'
             '其类内差异大、视觉特征分散，故检测精度较低（mAP@0.5 约 0.04），'
             '不纳入系统核心检测精度考核，仅用于对无法明确归类的拾得失物做兜底归类。')

# 4) save back to original filename
d.save(ORIG)
print('overwritten ORIG with 12-class table + note')

# 5) verify
d2 = docx.Document(ORIG)
t2 = d2.tables[2]
print('ORIG table[2] cols=', len(t2.columns), 'rows=', len(t2.rows))
for r in t2.rows:
    print(' | '.join(c.text.strip().replace('\n', ' ') for c in r.cells))
