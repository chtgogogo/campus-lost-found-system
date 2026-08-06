# -*- coding: utf-8 -*-
"""P4 论文补丁前期侦察：导出段落索引/样式/文本 + 表格内容，供定位修复点。"""
import sys
from pathlib import Path
from docx import Document

DOCX = Path(r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx")
OUT = Path(r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/tools/_p4_recon_out.txt")


def main():
    doc = Document(str(DOCX))
    lines = []
    lines.append(f"=== PARAGRAPHS ({len(doc.paragraphs)}) ===")
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "?"
        txt = p.text.replace("\n", " ").strip()
        lines.append(f"{i}\t[{style}]\t{txt[:160]}")
    lines.append(f"\n=== TABLES ({len(doc.tables)}) ===")
    for ti, t in enumerate(doc.tables):
        lines.append(f"-- table {ti}: {len(t.rows)} rows x {len(t.columns)} cols --")
        for ri, row in enumerate(t.rows):
            cells = [c.text.replace("\n", " ").strip()[:50] for c in row.cells]
            lines.append(f"  r{ri}\t" + " | ".join(cells))
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"written {OUT} | paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")


if __name__ == "__main__":
    main()
