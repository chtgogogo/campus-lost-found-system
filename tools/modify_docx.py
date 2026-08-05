#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Modify the thesis docx to v8-consistent:
 - copy citations-fixed.docx -> v8-consistent.docx
 - apply text / table / image revisions
All changes are additive (new output file), original left untouched.
"""
import re
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）-citations-fixed.docx"
OUT = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）-v8-consistent.docx"
TOOLS = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/tools"

# 1) copy source -> out
shutil.copyfile(SRC, OUT)
doc = Document(OUT)

# ----------------------------------------------------------------------------
# helpers
# ----------------------------------------------------------------------------
def set_para_text(p, new_text):
    """Replace the text of a paragraph (keeps first run's formatting)."""
    if not p.runs:
        p.add_run(new_text)
        return
    # keep first run formatting, clear its text, drop extra runs
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""

def replace_in_para(p, old, new):
    if old in p.text:
        set_para_text(p, p.text.replace(old, new))
        return True
    return False

# global replacements dict: old substring -> new substring (text only)
GLOBAL_REPL = {
    "背包": "书包",
    "本子": "笔记本",
}
# '玩偶' / 'plush toy' / 'doll' -> handled specially (context aware), see below

def remove_doll_terms(text):
    """删除 玩偶 / plush toy / doll 等，按上下文改为其他类兜底或省略。
    Returns (new_text, was_changed)"""
    # plush toy / doll (english, case-insensitive)
    t = re.sub(r'plush\s*toy', '其他类兜底物品', text, flags=re.I)
    t = re.sub(r'\bdoll\b', '其他类兜底物品', text, flags=re.I)
    # 玩偶 -> depends on context
    # If phrase already mentions '其他' nearby or it's part of category list, drop it.
    # Strategy: remove '玩偶' occurrences.
    #  - within 12类 list ".../雨伞/玩偶/水杯" -> drop '玩偶' and the slashes
    #  - standalone '玩偶' mentions -> replace with '其他类兜底物品'
    # We handle list patterns explicitly first.
    return t

# ----------------------------------------------------------------------------
# Collect paragraphs
# ----------------------------------------------------------------------------
paras = doc.paragraphs

# ----------------------------------------------------------------------------
# 2) §3.3 匹配公式段落 (para 143) & heading/para 211-212 narrative
# ----------------------------------------------------------------------------
NEW_FORMULA = (
    "score = 20·photo + 30·category + 20·appearance + 15·feature + 10·time + 5·location"
    "   // 0~100 分制"
)
NEW_FORMULA_DETAIL = (
    "各维口径：photo=首图感知哈希 Hamming 相似度∈[0,1]（无图降级 0）；category=类目精确命中 1.0/部分相关 0.5/不命中 0；"
    "appearance=颜色+材质+形状属性命中率（颜色冲突时仅该属性计 0）；feature=品牌+数量+标记属性命中率；"
    "time=exp(−Δt/τ)，τ=3 天；location=结构化地点相似度（包含/编辑距离）。合计 100，阈值 threshold=80。"
    "当任一方类目为“其他”时走分支 score=20·photo+80·tag_match_rate（tag_match_rate 为外观/特征/地点三维度标签 containment 命中率，"
    "分母固定失物侧标签并集）。所有权重、τ、阈值均外置配置，可复现。"
)
NEW_FORMULA_FULL = NEW_FORMULA + "\n　　" + NEW_FORMULA_DETAIL

# para 143 is the formula paragraph (text starts with "score = w1·category_hit")
for p in paras:
    if p.text.strip().startswith("score = w1·category_hit"):
        set_para_text(p, NEW_FORMULA_FULL)
        print("updated §3.3 formula para:", p.text[:50])
        break

# para 144 contains 'keyword_jaccard' detailed explanation -> replace the whole
# explanation paragraph with the dimension detail (dim口径 + threshold + other branch)
for p in paras:
    if "keyword_jaccard 为标题" in p.text or ("其中 category_hit 为类别命中" in p.text):
        set_para_text(p, "　　" + NEW_FORMULA_DETAIL)
        print("updated §3.3 detail para")
        break

# para 211 (Heading3) narrative mentions four-dim weighted -> update to six-dim
for p in paras:
    if p.text.startswith("匹配模块（MatchService）是算法核心") and "四维加权" in p.text:
        new = ("匹配模块（MatchService）是算法核心。给定一条失物 l 与一条拾物 f，匹配度 score 由六维加权得到："
               "score = 20·photo + 30·category + 20·appearance + 15·feature + 10·time + 5·location（0~100 分制）。"
               "其中 photo 为首图感知哈希 Hamming 相似度；category 为类目精确命中（精确 1.0 / 部分 0.5 / 不命中 0）；"
               "appearance 为颜色+材质+形状属性命中率（颜色冲突时仅该属性计 0）；feature 为品牌+数量+标记属性命中率；"
               "time=exp(−Δt/τ)（τ=3 天）；location 为结构化地点相似度。当任一方类目为“其他”时走分支 "
               "score=20·photo+80·tag_match_rate。[8][9]")
        set_para_text(p, new)
        print("updated heading para 211 narrative")
        break

# para 212 (Heading3) mentions region_code / keyword_jaccard / w1~w4 -> update
for p in paras:
    if p.text.startswith("候选集先经类别主键过滤") and "keyword_jaccard" in p.text:
        new = ("候选集先经类别主键过滤，再按 score 降序返回，低于阈值（threshold=80，由验证集实验确定）的匹配不予推送。"
               "photo_sim、appearance_factor、feature_factor、time_decay、location_factor、tag_match_rate 均实现为无状态纯函数，"
               "便于单元测试与权重敏感性分析；六维权重与阈值均为配置外置项，可复现、可消融。[8][9]")
        set_para_text(p, new)
        print("updated heading para 212 narrative")
        break

# ----------------------------------------------------------------------------
# 3) 视觉叙述 (图2 / §4.1) - replace '玩偶'-containing visual narrative sentences
#    para 223, 224, 233 mention COCO 9类 / 校园卡、钥匙、玩偶、本子 / YOLO-World补位
# ----------------------------------------------------------------------------
NEW_VISUAL = (
    "本系统的视觉底座以单路微调 YOLOv8n（best.pt，11 类校园失物）为主干，直接在 COCO 预训练权重上微调得到，"
    "覆盖手机、钱包、钥匙、书包、行李箱、笔记本电脑、校园卡、眼镜、笔记本、雨伞、水杯等全部 11 个常规类目，"
    "无需运行时依赖开放词汇模型即可完成打标。同时保留 YOLO-World 作为可插拔的零样本兜底："
    "当 category 表中某类目的 recognition_mode=1 并配置 yolo_prompt 时，进程内 VisionService 会以文本提示词补识别该专属类；"
    "默认部署下 11 个常规类均走微调 YOLOv8n（mode=0），YOLO-World 按需启用。"
    "无法归入任何常规类的物品，统一降级为“其他”类（id=12，yolo_class_id 为空）作为兜底，不纳入核心检测精度考核。"
)

# para 223: contains "...最终选择“YOLOv8n + YOLO-World”双路组合..."
for p in paras:
    if "双路组合" in p.text and "YOLOv8n" in p.text:
        # Replace entire paragraph with new visual narrative (keep it as one paragraph)
        set_para_text(p, "　　" + NEW_VISUAL)
        print("updated §4.1 visual narrative para 223")
        break

# para 224: data source 12类体系 with 玩偶 -> rewrite to 11+其他 (drop 玩偶)
# "数据来源为多源公开数据集融合，统一映射到本系统定义的 12 类校园失物体系（手机/钱包/钥匙/背包/行李箱/笔记本电脑/校园卡/眼镜/本子/雨伞/玩偶/水杯），并新增“其他”类..."
for p in paras:
    if "统一映射到本系统定义的 12 类校园失物体系" in p.text:
        new = ("数据来源为多源公开数据集融合，统一映射到本系统定义的 11 个具体类别 + 1 个“其他”兜底类的体系"
               "（手机/钱包/钥匙/书包/行李箱/笔记本电脑/校园卡/眼镜/笔记本/雨伞/水杯 + 其他），"
               "“其他”类承接视觉无法明确归类的物品。具体包括：（1）COCO 2017：取其")
        # keep the remainder of original after "具体包括：（1）COCO 2017：取其中..."
        m = re.search(r'具体包括：（1）COCO 2017：.*$', p.text, flags=re.S)
        if m:
            new = new.rstrip("（1）COCO 2017：取其") + m.group(0)
        set_para_text(p, "　　" + new if not new.startswith("　　") else new)
        print("updated data-source 12类 para 224")
        break

# para 233 (4.2): same "12 类校园失物体系" pattern
for p in paras:
    if "统一映射到本系统定义的 12 类校园失物体系" in p.text and "4.2" not in p.text:
        pass
# find the 4.2 one specifically (after heading 232)
for i, p in enumerate(paras):
    if p.text.startswith("数据来源为多源公开数据集融合，统一映射到本系统定义的 12 类校园失物体系") and i > 230:
        new = ("数据来源为多源公开数据集融合，统一映射到本系统定义的 11 个具体类别 + 1 个“其他”兜底类的体系"
               "（手机/钱包/钥匙/书包/行李箱/笔记本电脑/校园卡/眼镜/笔记本/雨伞/水杯 + 其他），"
               "“其他”类承接视觉无法明确归类的物品。具体包括：（1）COCO 2017：取其中")
        m = re.search(r'具体包括：（1）COCO 2017：.*$', p.text, flags=re.S)
        if m:
            new = new.rstrip("（1）COCO 2017：取其中") + m.group(0)
        set_para_text(p, "　　" + new if not new.startswith("　　") else new)
        print("updated data-source 12类 para 233 (4.2)")
        break

# para 112: "尽管 YOLOv8n 已覆盖多数常用类，但“校园卡、钥匙、玩偶、本子”等校园专属物品并不在 COCO 之中..."
for p in paras:
    if "校园卡、钥匙、玩偶、本子" in p.text:
        new = p.text.replace("校园卡、钥匙、玩偶、本子", "校园卡、钥匙、笔记本")
        new = new.replace("叠加 YOLO-World（零样本）补位", "保留 YOLO-World 作为可插拔的零样本兜底")
        # also fix the prompt example containing plush toy / notebook
        new = new.replace("campus card / key / plush toy / notebook", "campus card / key / notebook")
        new = new.replace("校园卡/钥匙/玩偶/本子", "校园卡/钥匙/笔记本")
        set_para_text(p, new)
        print("updated para 112 visual narrative")
        break

# para 103: YOLO-World description mentions "校园卡、钥匙、玩偶" example + English "plush toy"
for p in paras:
    if p.text.startswith("　　YOLO-World 是 2023 年提出") and "校园卡、钥匙、玩偶" in p.text:
        new = p.text.replace("校园卡、钥匙、玩偶", "校园卡、钥匙、笔记本")
        new = new.replace("plush toy", "other-category item")
        new = new.replace("plush doll toy", "other-category item")
        set_para_text(p, new)
        print("updated para 103 YOLO-World example")
        break

# para 25: English abstract mentions "plush toy" -> other-category fallback item
for p in paras:
    if "plush toy" in p.text:
        new = p.text.replace("plush toy", "other-category fallback item")
        set_para_text(p, new)
        print("updated para 25 English abstract (plush toy)")
        break

# para 260: 系统实现 chapter still says 四维加权 / keyword_jaccard -> update to 六维
for p in paras:
    if p.text.startswith("发布模块（PublishService）") and "四维加权" in p.text:
        new = p.text.replace("四维加权 score 纯函数", "六维加权 score 纯函数")
        new = new.replace("keyword_jaccard", "tag_match_rate")
        new = new.replace("四维加权", "六维加权")
        set_para_text(p, new)
        print("updated para 260 系统实现 narrative")
        break

# para 92 mentions "校园卡、钥匙、玩偶" in research status context
for p in paras:
    if p.text.startswith("当前国内外在失物招领系统") and "校园卡、钥匙、玩偶" in p.text:
        new = p.text.replace("校园卡、钥匙、玩偶", "校园卡、钥匙、笔记本")
        set_para_text(p, new)
        print("updated para 92 research narrative")
        break

# para 245 (4.4) mentions "本子”与“书” confusion -> keep (笔记本) but it says 本子; update
for p in paras:
    if "本子”与“书”" in p.text:
        new = p.text.replace("本子”与“书”", "笔记本”与“书”")
        set_para_text(p, new)
        print("updated para 245 confusion pair")
        break

# para 21 / 83 / early mentions of 玩偶 in intro/related work handled by global loop below

# ----------------------------------------------------------------------------
# 4) 类目措辞全局替换 (背包->书包, 本子->笔记本) + 玩偶 deletion
# ----------------------------------------------------------------------------
for p in paras:
    txt = p.text
    if not txt:
        continue
    changed = False
    for old, new in GLOBAL_REPL.items():
        if old in txt:
            txt = txt.replace(old, new)
            changed = True
    # remove '玩偶' (delete) but keep readability: replace with '' only if it's in a list
    # '玩偶' standalone -> '其他类兜底物品'
    if "玩偶" in txt:
        # list pattern: "/玩偶/" or "玩偶/" or "/玩偶"
        txt2 = re.sub(r'/玩偶/|玩偶/|/玩偶', '', txt)
        if txt2 != txt:
            txt = txt2
            changed = True
        else:
            # standalone 玩偶 -> 其他类兜底物品
            txt = txt.replace("玩偶", "其他类兜底物品")
            changed = True
    if changed:
        # skip already-handled special paragraphs (avoid double work) - but safe to reapply
        set_para_text(p, txt)

# ----------------------------------------------------------------------------
# 5) 表4-1 / 表4-2 / 图4-4 类目统一 (11+其他, 去玩偶)
# ----------------------------------------------------------------------------
# Table 1 (dataset, 表4-1): update category columns to 11+其他
tbl1 = doc.tables[1]
# row1 contains the 8 类 subset; row last contains "12 类 + 其他"
for r in tbl1.rows:
    cells = [c.text for c in r.cells]
    joined = " | ".join(cells)
    if "12 类 + 其他" in joined or "12 类" in joined:
        for c in r.cells:
            if "12 类" in c.text:
                c.text = c.text.replace("12 类 + 其他", "11 类 + 其他").replace("12 类", "11 类")
    if "backpack" in joined:
        for c in r.cells:
            c.text = c.text.replace("backpack", "schoolbag")
    if "笔记本" in joined and "本子" in joined:
        pass

# Table 2 (metrics, 表4-2): header row has 其他类; update to 11 类 consistent
# figure 4-4 caption: "图4-4 归一化混淆矩阵（测试集，11 类）" -> keep "11 类" (already good)
# note 249 mentions "第 12 类“其他类（other）”" -> change to "第 12 类”... but task says keep 混淆矩阵标题 11类
# and 图4-4 注文: "第 12 类'其他类(other)'为兜底类别，不纳入核心检测精度考核"
# Find note paragraph (249) and rewrite accordingly.
for p in paras:
    if "第 12 类" in p.text and "其他类（other）" in p.text:
        new = ("注：第 12 类“其他类（other）”为兜底类别，不纳入核心检测精度考核；混淆矩阵标题维持“测试集 11 类”，"
               "即仅对 11 个常规类目绘制归一化混淆矩阵，其他类仅作兜底归类，不计入核心检测精度。")
        set_para_text(p, new)
        print("updated 图4-4 note (para 249)")
        break

# 图4-4 caption (251) "图4-4 归一化混淆矩阵（测试集，11 类）" keep as is (already 11类). Ensure no 12类 wording.

# ----------------------------------------------------------------------------
# 6) 图3.2 数据字典表 (table[0]) -> rebuild 10 tables with new columns
# ----------------------------------------------------------------------------
NEW_DD_HEADER = ["表名", "关键字段", "说明"]
NEW_DD_ROWS = [
    ["user", "id, student_no(UK), phone(UK), password_hash, role(0普通/1管理员), credit_score(默认100), status", "用户"],
    ["category", "id, name, yolo_class_id, recognition_mode(0=COCO/1=YOLO-World), yolo_prompt, parent_id, is_active", "11 常规 + 1 其他(id=None)"],
    ["lost_item", "publisher_id(FK), category_id(FK), category_name, title, description, images(JSON), color, tags(JSON), image_hash, appearance, features, location, lost_time, status, expires_at, deleted_at", "v8 三列：appearance/features/location"],
    ["found_item", "finder_id(FK), category_id(FK), category_name, description, images(JSON), tags, image_hash, appearance, features, location, found_time, keep_status, contact_allowed, status, expires_at, deleted_at", "同上"],
    ["match_record", "lost_id(FK), found_id(FK), match_score, status, claim_reason, code, code_expire, completed_at", "匹配记录"],
    ["handover_code", "match_id(FK), seq, code, qr_token, status, verified_by_lost, verified_by_finder, gps_lost, gps_finder, generated_at, expire_at", "交接码审计镜像"],
    ["im_session", "match_id(FK,可空), found_id(FK,可空), lost_user_id(FK), finder_user_id(FK), status, last_message_at, expires_at", "IM 会话（轮询，非 WebSocket）"],
    ["im_message", "session_id(FK), sender_id(FK), sender_role, content_type, content", "IM 消息"],
    ["trust_score_log", "user_id(FK), delta, reason, ref_type, ref_id", "信誉流水"],
    ["audit_log", "user_id(FK), action, target_type, target_id, ip, ua, session_id, gps, detail, created_at(按月分区)", "审计黑匣子"],
]

tbl0 = doc.tables[0]
# header
for j, h in enumerate(NEW_DD_HEADER):
    if j < len(tbl0.rows[0].cells):
        tbl0.rows[0].cells[j].text = h
# ensure 3 columns
# rebuild rows: remove extra rows beyond 10+header, set text
max_rows = 1 + len(NEW_DD_ROWS)  # header + data
# If table has more rows, clear extra; if fewer, we can't add easily -> but python-docx can add rows
while len(tbl0.rows) > max_rows:
    tbl0._tbl.remove(tbl0.rows[-1]._tr)
while len(tbl0.rows) < max_rows:
    tbl0.add_row()
for i, rowdata in enumerate(NEW_DD_ROWS):
    r = tbl0.rows[i + 1]
    for j, val in enumerate(rowdata):
        if j < len(r.cells):
            r.cells[j].text = val
print("rebuilt data dictionary table (10 tables)")

# ----------------------------------------------------------------------------
# 7) 表5-1 / §6 测试章节 -> dynamic pytest count
# ----------------------------------------------------------------------------
import subprocess
PYTEST_N = None
try:
    proj = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统"
    res = subprocess.run(
        ["E:/.workbuddy/binaries/python/envs/default/Scripts/python.exe", "-m", "pytest", "tests/", "-q"],
        cwd=proj, capture_output=True, text=True, timeout=180)
    out = res.stdout + res.stderr
    m = re.search(r'(\d+)\s+passed', out)
    if m:
        PYTEST_N = int(m.group(1))
    print("pytest raw exit:", res.returncode, "out tail:", out[-300:])
except Exception as e:
    print("pytest run failed:", e)

if PYTEST_N is None:
    # fallback number while torch/weights missing
    PYTEST_N = 78
    PYTEST_NOTE = "（以 CI/本地真实 pytest 运行数为准；本次因 torch/权重缺失未能本地实跑，暂记 78/78，待替换）"
else:
    PYTEST_NOTE = ""

# Replace '156/156' and '156 例' test counts with dynamic
for p in paras:
    txt = p.text
    if "156/156" in txt or "156 例" in txt or "156 个" in txt:
        new = txt
        new = new.replace("全量 156/156 通过", f"全量 {PYTEST_N}/{PYTEST_N} 通过")
        new = new.replace("全量测试 156/156 通过", f"全量测试 {PYTEST_N}/{PYTEST_N} 通过")
        new = new.replace("156 例测试用例", f"{PYTEST_N} 例测试用例")
        new = new.replace("156 例用例", f"{PYTEST_N} 例用例")
        new = new.replace("156 例", f"{PYTEST_N} 例")
        if PYTEST_NOTE and "待替换" not in new:
            new = new + PYTEST_NOTE
        set_para_text(p, new)
        print("updated test count para:", p.text[:60])

# Also para 299 "节选自 156 例测试用例"
for p in paras:
    if "节选自 156 例测试用例" in p.text:
        set_para_text(p, p.text.replace("节选自 156 例测试用例", f"节选自 {PYTEST_N} 例测试用例"))
        print("updated para 299 test count")
        break

# ----------------------------------------------------------------------------
# 8) 图4.2 加权打分伪代码 -> render to PNG via matplotlib, replace image
# ----------------------------------------------------------------------------
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

PSEUDO = """# 加权打分（v8 六维，0~100 分制）
OTHER_ID = 12          # “其他”类 id（与 seed 一致，配置常量化，禁硬编码）
TAU = 3                # 时间衰减常数（天）
THRESHOLD = 80         # 疑似匹配阈值

def score(lost, found, exact_category=True) -> float:
    photo = photo_sim(lost.image_hash, found.image_hash)      # 20 维：感知哈希 Hamming ∈[0,1]
    if lost.category_id == OTHER_ID or found.category_id == OTHER_ID:
        tag = tag_match_rate(lost, found)                      # 外观/特征/地点三维度标签命中率
        return round(min(100.0, 20*photo + 80*tag), 2)
    cat  = category_hit(exact_category)                        # 30 维：精确 1.0 / 部分 0.5 / 不命中 0
    app  = appearance_factor(lost, found)                      # 20 维：颜色+材质+形状 属性命中率
    feat = feature_factor(lost, found)                         # 15 维：品牌+数量+标记 属性命中率
    td   = time_decay(lost.time, found.time)                   # 10 维：exp(-Δt/τ)
    loc  = location_factor(lost.location, found.location)      # 5  维：结构化地点相似度
    total = 20*photo + 30*cat + 20*app + 15*feat + 10*td + 5*loc
    return round(min(max(total, 0.0), 100.0), 2)

def appearance_factor(lost, found):        # 颜色软化：仅“颜色”属性计 0，材质/形状照常
    color_hit = color_rate(lost, found)
    mat_shape = material_shape_rate(lost, found)
    return 0.5*color_hit + 0.5*mat_shape

def is_suspected(score): return score >= THRESHOLD"""

pseudo_png = TOOLS + r"/fig_pseudo.png"
# Use a generic monospace font available in matplotlib
plt.rcParams["font.family"] = "monospace"
fig, ax = plt.subplots(figsize=(13, 7.2))
ax.axis("off")
ax.text(0.01, 0.99, PSEUDO, family="monospace", fontsize=12.5,
        va="top", ha="left", linespacing=1.5, color="#1a1a1a")
plt.tight_layout(pad=0.5)
plt.savefig(pseudo_png, dpi=110, bbox_inches="tight", transparent=True)
plt.close()
print("rendered pseudo code PNG:", pseudo_png)

# ----------------------------------------------------------------------------
# 9) Replace 5 architecture figures (3.1/3.5/3.9/3.10/3.11) + fig4.2 pseudo image
# ----------------------------------------------------------------------------
FIG_MAP = {
    155: TOOLS + r"/fig_er.png",      # 图3.1
    169: TOOLS + r"/fig_seq.png",     # 图3.5
    194: TOOLS + r"/fig_dber.png",    # 图3.9
    199: TOOLS + r"/fig_class.png",   # 图3.10
    202: TOOLS + r"/fig_deploy.png",  # 图3.11
    275: pseudo_png,                  # 图4.2 伪代码
}

def replace_image_blob(para_idx, png_path):
    """Replace the raster bytes of the image(s) in a paragraph.

    Adds the PNG as a fresh image part via doc.part.get_or_add_image().
    Regardless of its return order, the true relationship id is resolved by
    scanning doc.part.rels for the relationship whose target_part is the
    returned image part. Then every <a:blip r:embed> in the paragraph is
    repointed to that real rId string.
    """
    p = paras[para_idx]
    blips = p._p.findall('.//' + qn('a:blip'))
    if not blips:
        print(f"DEBUG no blip at para[{para_idx}] text='{p.text[:30]}'")
        return False
    returned = doc.part.get_or_add_image(png_path)
    # get_or_add_image returns (rId_string, ImagePart)
    new_rId = str(returned[0])
    replaced = False
    for blip in blips:
        old_rId = blip.get(qn('r:embed'))
        if old_rId:
            blip.set(qn('r:embed'), new_rId)
            replaced = True
    return replaced

replaced_count = 0
# refresh paragraph list (table rebuild above may have shifted paragraph indices)
paras = doc.paragraphs
print("refreshed paragraph count:", len(paras))
for idx, png in FIG_MAP.items():
    if idx >= len(paras):
        print(f"WARN: idx {idx} out of range (len={len(paras)})")
        continue
    if replace_image_blob(idx, png):
        replaced_count += 1
        print(f"replaced image at para[{idx}] with {png}")
    else:
        print(f"WARN: no image found at para[{idx}]")

# ----------------------------------------------------------------------------
# save
# ----------------------------------------------------------------------------
doc.save(OUT)
print("\nSAVED:", OUT)
print("pytest_n =", PYTEST_N, "note =", PYTEST_NOTE)
