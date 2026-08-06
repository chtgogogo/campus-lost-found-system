# -*- coding: utf-8 -*-
"""
论文同步 第三批补丁（收尾）：
  1. 剩余 4 处「已部署骨干」语境 YOLOv8n -> YOLOv8s（[103][228][230][235]）
  2. 数据划分 7:2:1 -> 真实张数（[225][234] + 表格 1 末行）
  3. 修复原文重复错字「具体包括具体包括」
  4. 表格 1 合并行样本规模填真实数字
"""
import os
import shutil
import datetime as _dt
from docx import Document

THESIS = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
BAK = THESIS + ".bak.p3." + _dt.datetime.now().strftime("%Y%m%d_%H%M%S")

N_TRAIN, N_VAL, N_TEST = 59309, 18318, 3236
TOTAL = N_TRAIN + N_VAL + N_TEST
R = (N_TRAIN / TOTAL * 100, N_VAL / TOTAL * 100, N_TEST / TOTAL * 100)
SPLIT_TXT = ("三源按类别分层抽样合并后共 %s 张图像，划分为训练集/验证集/测试集 = %s / %s / %s 张（约 %.0f%%:%.0f%%:%.0f%%）"
             % (format(TOTAL, ","), format(N_TRAIN, ","), format(N_VAL, ","),
                format(N_TEST, ","), *R))

doc = Document(THESIS)
if not os.path.exists(BAK):
    shutil.copy(THESIS, BAK)
    print("backup ->", BAK)


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def rep(idx, old, new):
    p = doc.paragraphs[idx]
    assert old in p.text, "[%d] 未找到: %r\n实际: %s" % (idx, old, p.text[:150])
    set_text(p, p.text.replace(old, new))


# ---- 1. 部署语境骨干名 ----
rep(103, "二者结合——YOLOv8n 处理 COCO 重叠常用类",
    "二者结合——微调 YOLOv8s 处理 COCO 重叠常用类")
rep(228, "本章明确了以 YOLOv8n + YOLO-World 双路检测作为视觉底座",
    "本章明确了以微调 YOLOv8s + YOLO-World 双路检测作为视觉底座")
rep(230, "本文以 YOLOv8n（COCO 预训练）叠加 YOLO-World（零样本）作为视觉底座",
    "本文以 YOLOv8s（COCO 预训练权重微调）叠加 YOLO-World（零样本）作为视觉底座"
    "（选型初期以 YOLOv8n 为基线，因稀有类精度不足升级至 s 档，对比见 4.4 节）")
rep(235, "尺寸归一化至模型输入（YOLOv8n img_size=640）",
    "尺寸归一化至模型输入（YOLOv8s img_size=640）")

# ---- 2. 数据划分 + 重复错字（两处同文段落） ----
for i in (225, 234):
    rep(i, "具体包括具体包括：", "具体包括：")
    rep(i, "三源按类别分层抽样合并为训练集/验证集/测试集 = 7:2:1；", SPLIT_TXT + "；")

# ---- 3. 表格 1 末行 ----
t = doc.tables[1]
last = t.rows[-1]
assert "7:2:1" in last.cells[3].text, "表格1 末行样本规模列内容异常: %r" % last.cells[3].text
last.cells[3].text = "%s 图（train %s / val %s / test %s）" % (
    format(TOTAL, ","), format(N_TRAIN, ","), format(N_VAL, ","), format(N_TEST, ","))
last.cells[1].text = "11 类 + 其他（nc=12）"

doc.save(THESIS)
print("P3 DONE ->", THESIS)
print(" split:", SPLIT_TXT)
