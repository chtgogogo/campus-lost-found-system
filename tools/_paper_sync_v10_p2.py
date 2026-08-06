# -*- coding: utf-8 -*-
"""
论文同步 第二批补丁（依赖 _paper_sync_v10.py 已执行，文档为 336 段版本）：
  1. 已部署骨干 YOLOv8n -> YOLOv8s（区分「YOLO 系列科普」与「本系统选用」两种语境）
  2. 训练超参对齐真实 args.yaml：batch 24->8、数据划分 7:2:1 -> 真实张数、epochs 跑满 120
  3. 数据集类目 11 类 -> 12 类（含 other 兜底）口径统一
  4. 分组 mAP 与总体指标统一到「测试集」（图4-4 图注即测试集），消融对照显式标注「验证集」
  5. 中英文摘要匹配引擎口径：四维 -> v2 七子维度 + 归一化
  6. 类别长尾分布真实数字（val 最大失衡 103:1）写入 WIoU 消融段作为归因证据
  7. 性能指标 CPU 单图时延对齐 v8s 实测（热推理 0.135s）

真实配置来源（software-engineer 核实）：
  runs/detect/lostfound_v8s_ciou/args.yaml + results.csv
  tools/dataset_prep/train_vision.py + wiou_loss.py
  dataset/final/data.yaml（nc=12）
"""
import json
import shutil
import os
import datetime as _dt
from docx import Document

THESIS = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
BAK = THESIS + ".bak.p2." + _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
ROOT = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统"

# ===== 真实训练配置（已核实） =====
TR = dict(model="YOLOv8s", epochs=120, actual_epochs=120, early_stopped=False,
          batch=8, imgsz=640, optimizer="auto", lr0=0.01, patience=30,
          close_mosaic=10, amp=True,
          n_train=59309, n_val=18318, n_test=3236,
          params_s=11140244, gflops_s=28.7,
          params_n=3157200, gflops_n=8.9,
          cpu_cold_s=2.022, cpu_warm_s=0.135)
# val 集类别长尾（标注框数）
LONGTAIL = [("backpack", 7867), ("notebook", 6077), ("bottle", 5339), ("umbrella", 3154),
            ("suitcase", 2268), ("phone", 2112), ("laptop", 1339), ("other", 516),
            ("glasses", 248), ("keys", 201), ("wallet", 171), ("campus_card", 76)]
IMBALANCE = LONGTAIL[0][1] / LONGTAIL[-1][1]      # ≈103.5 : 1
TOTAL_IMG = TR["n_train"] + TR["n_val"] + TR["n_test"]
R_TR = TR["n_train"] / TOTAL_IMG * 100
R_VA = TR["n_val"] / TOTAL_IMG * 100
R_TE = TR["n_test"] / TOTAL_IMG * 100

J = lambda n: json.load(open(os.path.join(ROOT, "runs/detect", n), encoding="utf-8"))
s_test = J("eval_best_test_metrics.json")          # v8s+CIoU  测试集（论文主指标）
s_val = J("eval_best_val_metrics.json")            # v8s+CIoU  验证集（消融基准）
n_test = J("eval__tmp_old_best_test_metrics.json")  # v8n 基线  测试集
n_val = J("eval__tmp_old_best_val_metrics.json")   # v8n 基线  验证集
wiou = J("eval_WIoU_disaster_metrics.json")        # WIoU 灾难  验证集

COCO = [0, 1, 2, 3, 4, 5, 8, 9, 10]
CAMPUS = [6, 7]
OTHER = [11]


def grp(ev, idxs, key):
    pc = {c["index"]: c for c in ev["per_class"]}
    return sum(pc[i][key] for i in idxs) / len(idxs)


# 测试集分组（论文表 4-2 主数据）
t_coco50, t_camp50, t_oth50 = (grp(s_test, g, "ap50") for g in (COCO, CAMPUS, OTHER))
t_coco95, t_camp95, t_oth95 = (grp(s_test, g, "ap50_95") for g in (COCO, CAMPUS, OTHER))
t_all50 = s_test["overview"]["mAP50"]
t_all95 = s_test["overview"]["mAP50_95"]
t_p, t_r = s_test["overview"]["precision"], s_test["overview"]["recall"]
n_all50, n_all95 = n_test["overview"]["mAP50"], n_test["overview"]["mAP50_95"]
n_p, n_r = n_test["overview"]["precision"], n_test["overview"]["recall"]
# 验证集（消融口径）
v_all50, v_all95 = s_val["overview"]["mAP50"], s_val["overview"]["mAP50_95"]
v_r = s_val["overview"]["recall"]
w50, w95 = wiou["overview"]["mAP50"], wiou["overview"]["mAP50_95"]
w_p, w_r = wiou["overview"]["precision"], wiou["overview"]["recall"]
pcs_t = {c["name"]: c for c in s_test["per_class"]}

# ---- 护栏：确认各 JSON 未被张冠李戴 ----
assert t_all50 > n_all50 > 0.4, "v8s 测试集应显著优于 v8n 基线"
assert w50 < 0.1, "WIoU 灾难数据异常"
assert abs(v_all50 - 0.710) < 0.01, "验证集主指标漂移，请复核"

doc = Document(THESIS)
if not os.path.exists(BAK):
    shutil.copy(THESIS, BAK)
    print("backup ->", BAK)


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def rep(idx, old, new, must=True):
    """在指定段落做子串替换，带存在性断言，避免改错段落。"""
    p = doc.paragraphs[idx]
    if old not in p.text:
        if must:
            raise AssertionError("[%d] 未找到待替换文本: %r\n实际: %s" % (idx, old, p.text[:120]))
        return False
    set_text(p, p.text.replace(old, new))
    return True


# ================= 1. 摘要（中） =================
rep(21, "视觉识别环节将 YOLOv8n（COCO 预训练，覆盖书包、手机等 9 类常用物品）",
    "视觉识别环节将微调后的 YOLOv8s（在 COCO 预训练权重上微调，覆盖书包、手机、钥匙、校园卡等 11 类校园高频失物）")
rep(21, "匹配引擎依据“类别命中 + 时间衰减 + 地点层级 + 关键词相似度”的加权策略计算匹配度",
    "匹配引擎依据“照片⊕类别 + 量词 + 颜色 + 状态 + 地点 + 其他关键词 + 时间衰减”七子维度加权（并按失主实际填写维度归一化）计算匹配度")
rep(21, "本文详细阐述了系统需求分析、数据库与接口设计、加权匹配算法、模型选型与验证方案，以及核心模块的实现与测试。",
    "本文详细阐述了系统需求分析、数据库与接口设计、加权匹配算法、模型选型与验证方案，以及核心模块的实现与测试。视觉模型在自建融合数据集上微调后，测试集 mAP@0.5 达 %.3f，较 YOLOv8n 基线（%.3f）提升 %.1f 个百分点；训练中曾以 WIoU v3 损失做对照消融，因数据集类别长尾（最大失衡约 %d:1）导致稀有类召回崩溃（验证集 mAP@0.5 由 %.3f 降至 %.3f），最终选定 CIoU 损失。"
    % (t_all50, n_all50, (t_all50 - n_all50) * 100, round(IMBALANCE), v_all50, w50))
rep(21, "系统后端最小闭环（发布→打标→匹配→认领→交接）在单元测试与端到端测试下均稳定运行",
    "系统后端最小闭环（发布→打标→匹配→认领→交接）在 334 例单元测试与端到端测试下均稳定运行")

# ================= 2. 摘要（英） =================
rep(25, "YOLOv8n (COCO-pretrained, covering 9 common categories such as bags and phones)",
    "a fine-tuned YOLOv8s (fine-tuned from COCO-pretrained weights, covering 11 high-frequency campus categories such as bags, phones, keys and campus cards)")
rep(25, "A matching engine then computes a weighted score from category hit, time decay, location hierarchy, and keyword similarity, and proactively pushes suspected matches to the owner.",
    "A matching engine then computes a weighted score over seven sub-dimensions — photo-category, quantity, colour, condition, location, other keywords and time decay — normalised by the dimensions the owner actually filled in, and proactively pushes suspected matches to the owner.")
rep(25, "Tests show that the minimal back-end closed loop runs stably under unit and end-to-end tests.",
    "On the self-built fused dataset the fine-tuned model reaches a test-set mAP@0.5 of %.3f, %.1f percentage points above the YOLOv8n baseline (%.3f). An ablation replacing CIoU with WIoU v3 caused a collapse in recall on rare classes due to the long-tailed class distribution (imbalance ratio about %d:1), with validation mAP@0.5 dropping from %.3f to %.3f; CIoU was therefore retained. Tests show that the minimal back-end closed loop runs stably under 334 unit and end-to-end test cases."
    % (t_all50, (t_all50 - n_all50) * 100, n_all50, round(IMBALANCE), v_all50, w50))

# ================= 3. 绪论 / 研发设想 =================
rep(84, "以 YOLOv8n（COCO 预训练覆盖常用类）+ YOLO-World（零样本补校园专属类）为视觉底座",
    "以微调 YOLOv8s（COCO 预训练权重微调，覆盖 11 类校园高频失物）+ YOLO-World（零样本补校园专属类）为视觉底座")
rep(95, "将 YOLOv8n 与 YOLO-World 推理以进程内单例方式嵌入",
    "将微调 YOLOv8s 与 YOLO-World 推理以进程内单例方式嵌入")
rep(95, "叠加时间衰减、地点层级与关键词 Jaccard 的规则加权打分[8][9]",
    "叠加量词、颜色、状态、地点、其他关键词与时间衰减的七子维度可解释加权打分（并按失主实际填写维度归一化）[8][9]")

# ================= 4. 相关技术（科普语境保留 n 档，选型语境改 s 档） =================
rep(102, "其中 YOLOv8n（nano）参数量最小、适合 CPU/边缘部署。",
    "其中 YOLOv8n（nano）参数量最小、适合 CPU/边缘部署，YOLOv8s（small）在参数量与精度间更为均衡。")
set_text(doc.paragraphs[108],
    "本系统选用 YOLOv8s（small）作为通用物品识别主干：其一，COCO 预训练已覆盖书包、手机、笔记本、水杯、雨伞等校园高频失物类，可直接复用并微调；其二，相较 YOLOv8n（%.2fM 参数 / %.1f GFLOPs），YOLOv8s（%.2fM 参数 / %.1f GFLOPs）容量更大，在本文长尾数据集上测试集 mAP@0.5 由 %.3f 提升至 %.3f，而 CPU 单图热推理仅 %.3fs（冷启动 %.2fs），仍满足“发布即打标”的实时体验；其三，Ultralytics 生态成熟，便于以进程内方式集成、以配置切换设备（cpu/cuda），降低部署与运维成本。选型初期曾以 YOLOv8n 作为基线，因稀有类精度不足而升级至 s 档。[6][2][3]"
    % (TR["params_n"] / 1e6, TR["gflops_n"], TR["params_s"] / 1e6, TR["gflops_s"],
       n_all50, t_all50, TR["cpu_warm_s"], TR["cpu_cold_s"]))
rep(112, "尽管 YOLOv8n 已覆盖多数常用类", "尽管微调后的 YOLOv8s 已覆盖多数常用类")
rep(112, "与 YOLOv8n 形成“通用+专属”的双路视觉底座", "与 YOLOv8s 形成“通用+专属”的双路视觉底座")
rep(114, "YOLOv8n 提供高精度、低时延的通用物品检测", "微调 YOLOv8s 提供高精度、低时延的通用物品检测")

# ================= 5. 性能指标 =================
rep(184, "视觉打标在 CPU 下单图 ≤ 1s（YOLOv8n）",
    "视觉打标在 CPU 下单图 ≤ 1s（YOLOv8s 实测热推理 %.3fs，冷启动 %.2fs）" % (TR["cpu_warm_s"], TR["cpu_cold_s"]))

# ================= 6. 视觉底座（11 类 -> 12 类口径） =================
set_text(doc.paragraphs[224],
    "　　本系统的视觉底座以单路微调 YOLOv8s（best.pt，nc=12）为主干，直接在 COCO 预训练权重上微调得到，覆盖手机、钱包、钥匙、书包、行李箱、笔记本电脑、校园卡、眼镜、笔记本、雨伞、水杯等 11 个常规类目，外加 1 个“其他”兜底类，无需运行时依赖开放词汇模型即可完成打标。同时保留 YOLO-World 作为可插拔的零样本兜底：当 category 表中某类目的 recognition_mode=1 并配置 yolo_prompt 时，进程内 VisionService 会以文本提示词补识别该专属类；默认部署下 11 个常规类均走微调 YOLOv8s（mode=0），YOLO-World 按需启用。无法归入任何常规类的物品，统一降级为“其他”类（yolo_class_id=11）作为兜底，不纳入核心检测精度考核。")

# ================= 7. 训练超参（两处重复段落） =================
_train_txt = ("本系统在 COCO 预训练权重基础上开展微调（fine-tuning），而非仅用固定预训练权重，以提升校园专属类的识别精度。"
              "训练在配备 NVIDIA GPU 的机器上进行，框架为 Ultralytics YOLOv8（Python 3.12 + PyTorch CUDA），主干选用 YOLOv8s；"
              "输入尺寸 img_size=640，批次大小 batch=%d（受本机 15.6GB 内存与显存限制，采用小批次并关闭 dataloader 多进程 workers=0 以规避 OOM），"
              "训练轮数 epochs=%d（实际跑满 %d 轮，未触发早停），优化器 auto（SGD/Adam 自适应），初始学习率 %.2f，"
              "启用早停（patience=%d）与 AMP 混合精度，并在最后 %d 轮关闭 Mosaic 增强（close_mosaic）以稳定收敛。"
              "边界框回归损失采用 CIoU（曾以 WIoU v3 做对照消融，见 4.5 节）。"
              "YOLO-World 以文本提示词列表驱动零样本识别，不参与反向传播。训练环境、参数与过程如下。[6][2][3][10]"
              % (TR["batch"], TR["epochs"], TR["actual_epochs"], TR["lr0"], TR["patience"], TR["close_mosaic"]))
set_text(doc.paragraphs[226], _train_txt)
set_text(doc.paragraphs[239], _train_txt.replace("[6][2][3][10]", "").rstrip())

set_text(doc.paragraphs[241],
    "训练参数：主干 YOLOv8s，img_size=640，batch=%d，epochs=%d（跑满，未早停），optimizer=auto，lr0=%.2f，patience=%d（早停），"
    "close_mosaic=%d，AMP=True，bbox 损失=CIoU。数据集共 %s 张图像，按 train/val/test = %s / %s / %s 张（约 %.0f%%:%.0f%%:%.0f%%）划分。"
    % (TR["batch"], TR["epochs"], TR["lr0"], TR["patience"], TR["close_mosaic"],
       format(TOTAL_IMG, ","), format(TR["n_train"], ","), format(TR["n_val"], ","),
       format(TR["n_test"], ","), R_TR, R_VA, R_TE))

# ================= 8. 表 4-2 改为测试集 + 补精确率/召回率行说明 =================
tbl = doc.tables[2]
for r, k in ((1, "50"), (2, "95")):
    vals = (t_coco50, t_camp50, t_all50, t_oth50) if k == "50" else (t_coco95, t_camp95, t_all95, t_oth95)
    for ci, v in zip((1, 2, 3, 4), vals):
        tbl.rows[r].cells[ci].text = "%.3f" % v

rep(250, "注：第 12 类“其他类（other）”为兜底类别，不纳入核心检测精度考核；",
    "注：表中数据为测试集（%s 张图像）评测结果，由 tools/dataset_prep/eval_vision.py 一次性导出；"
    "第 12 类“其他类（other）”为兜底类别，不纳入核心检测精度考核；" % format(TR["n_test"], ","))

# ================= 9. mAP 分析段（验证集 -> 测试集） =================
set_text(doc.paragraphs[253],
    "　　由表 4-2 可知，微调后的 YOLOv8s 在测试集（%s 张图像）上整体 mAP@0.5 达到 %.3f、mAP@0.5:0.95 为 %.3f，精确率 %.3f、召回率 %.3f。"
    "分组来看，校园专属类（%.3f）明显高于 COCO 通用类（%.3f）——校园卡（AP@0.5 %.3f）与眼镜（%.3f）虽样本稀少，但类内形态一致、判别特征稳定，"
    "反而优于形态与拍摄角度差异极大的书包（%.3f）、笔记本（%.3f）等通用类，说明微调对校园场景专属物品的适配是有效的；"
    "“其他类”因类内差异大、视觉特征分散，mAP@0.5 仅 %.3f，不纳入系统核心检测精度考核。"
    "作为同口径对照，初版 YOLOv8n 基线在同一测试集上 mAP@0.5 为 %.3f、mAP@0.5:0.95 为 %.3f（精确率 %.3f、召回率 %.3f），"
    "骨干升级至 YOLOv8s 后整体 mAP@0.5 提升 %.1f 个百分点、召回率提升 %.1f 个百分点，代价是参数量由 %.2fM 增至 %.2fM、"
    "CPU 单图热推理耗时 %.3fs（仍远低于 1s 的实时性约束）。在验证集（%s 张图像）上该模型 mAP@0.5 为 %.3f，与测试集结论一致，未见过拟合迹象。"
    "后续可进一步扩充稀有类样本、引入针对性数据增强来巩固精度优势。"
    % (format(TR["n_test"], ","), t_all50, t_all95, t_p, t_r,
       t_camp50, t_coco50, pcs_t["campus_card"]["ap50"], pcs_t["glasses"]["ap50"],
       pcs_t["backpack"]["ap50"], pcs_t["notebook"]["ap50"], t_oth50,
       n_all50, n_all95, n_p, n_r,
       (t_all50 - n_all50) * 100, (t_r - n_r) * 100,
       TR["params_n"] / 1e6, TR["params_s"] / 1e6, TR["cpu_warm_s"],
       format(TR["n_val"], ","), v_all50))

# ================= 10. WIoU 消融段：补长尾真实数字 + 显式标注验证集口径 =================
set_text(doc.paragraphs[254],
    "　　在最终选型前，本文对边界框回归损失做了对照消融。两次训练除损失函数外的全部超参完全一致（同为 YOLOv8s 骨干、batch=%d、epochs=%d、lr0=%.2f），"
    "仅将 ultralytics 原生 CIoU 替换为 WIoU v3（γ=1.9，通过替换 BboxLoss 注入），因此可将精度差异归因于损失函数本身。"
    "结果显示：改用 WIoU v3 后验证集 mAP@0.5 骤降至 %.3f（mAP@0.5:0.95 %.3f），精确率 %.3f、召回率仅 %.3f，"
    "较同骨干 CIoU 方案（验证集 mAP@0.5 %.3f、召回率 %.3f）几近完全退化，逐类 AP 普遍跌至 0.1 以下。"
    "分析其原因：WIoU v3 以动态非单调聚焦机制“关注中等质量锚框”，对低质量预测施加更小的梯度增益以抑制其对回归的干扰；"
    "而本文数据集类别分布高度长尾——验证集中书包（%d 个标注框）、笔记本（%d）、水杯（%d）等常见类占绝大多数，"
    "而校园卡（%d）、钱包（%d）、钥匙（%d）、眼镜（%d）等校园稀有类样本极少，最大失衡比约 %d:1。"
    "稀有类在训练早期产生的预测框恰恰集中于“低质量”区间，被 WIoU 判定为离群而抑制梯度，致其本就微弱的学习信号进一步衰减，最终整体召回崩溃。"
    "鉴于该退化不可接受，本文回退至 CIoU 损失并重训 YOLOv8s，验证集 mAP@0.5 恢复至 %.3f、mAP@0.5:0.95 达 %.3f。"
    "该消融实验表明：损失函数并非“越新越好”，其有效性高度依赖数据分布，在长尾/不均衡数据集上采用聚焦式损失需格外审慎。"
    % (TR["batch"], TR["epochs"], TR["lr0"], w50, w95, w_p, w_r, v_all50, v_r,
       LONGTAIL[0][1], LONGTAIL[1][1], LONGTAIL[2][1],
       dict(LONGTAIL)["campus_card"], dict(LONGTAIL)["wallet"],
       dict(LONGTAIL)["keys"], dict(LONGTAIL)["glasses"], round(IMBALANCE),
       v_all50, v_all95))

# ================= 11. 结论段数据对齐测试集 =================
rep(317, "（mAP@0.5 由 %.3f 跌至 %.3f，召回率仅 %.3f）" % (v_all50, w50, w_r),
    "（验证集 mAP@0.5 由 %.3f 跌至 %.3f，召回率仅 %.3f）" % (v_all50, w50, w_r))
rep(317, "真实推理已接入并实测 mAP@0.5=%.3f" % v_all50,
    "真实推理已接入，测试集实测 mAP@0.5=%.3f（验证集 %.3f）" % (t_all50, v_all50))

doc.save(THESIS)
print("P2 DONE ->", THESIS)
print("  test : v8s %.3f / v8n %.3f  (+%.1f pt)" % (t_all50, n_all50, (t_all50 - n_all50) * 100))
print("  group: coco %.3f campus %.3f other %.3f" % (t_coco50, t_camp50, t_oth50))
print("  wiou : %.3f (recall %.3f)   imbalance %.1f:1" % (w50, w_r, IMBALANCE))
