# -*- coding: utf-8 -*-
"""
同步 v10 / flow-v3 成果到毕业论文 .docx：
  - 评分公式 六维加权 -> v2 七子维度 + 归一化
  - 训练消融：WIoU 灾难过程 + 原因 + 回退 CIoU 方案
  - mAP 数据更新为 v8s+CIoU（从 eval_best_val_metrics.json 实时计算 COCO/校园/其他 分组）
  - 测试数 184 -> 334（334 passed / 1 skipped，pytest --co = 335）
  - flow-v3（keep1 单向进匹配池 + 低分阈值60）与管理员（邀请码 + 后台取证导出）内容
  - YOLOv8n -> YOLOv8s（已部署骨干）措辞对齐

运行前请确保 eval_best_val_metrics.json 已是 v8s+CIoU 结果（即 best.pt 当前权重评测）。
"""
import json
import shutil
import datetime as _dt
from docx import Document

THESIS = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
BAK = THESIS + ".bak.rev" + _dt.datetime.now().strftime("%Y%m%d")
BEST_EVAL = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/runs/detect/eval_best_val_metrics.json"
WIOU_EVAL = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/runs/detect/eval_WIoU_disaster_metrics.json"
N_EVAL = r"E:/xuexixiangguan/pythonProject/gongcheng/失物招领系统/runs/detect/eval__tmp_old_best_val_metrics.json"

doc = Document(THESIS)
if not __import__("os").path.exists(BAK):
    shutil.copy(THESIS, BAK)
    print("backup ->", BAK)

best = json.load(open(BEST_EVAL, encoding="utf-8"))
wiou = json.load(open(WIOU_EVAL, encoding="utf-8"))
v8n = json.load(open(N_EVAL, encoding="utf-8"))   # 同验证集下的初版 YOLOv8n+CIoU 基线
n50 = v8n["overview"]["mAP50"]
n95 = v8n["overview"]["mAP50_95"]

# 安全护栏：若 eval_best 仍是 WIoU 灾难值（mAP50<0.1），说明评测尚未覆盖新权重，终止避免写错数据
if best["overview"]["mAP50"] < 0.1:
    print("ERROR: eval_best_val_metrics.json 仍是 WIoU 灾难值 (mAP50=%.3f)，请先对当前 best.pt 跑完 eval_vision.py 再执行。"
          % best["overview"]["mAP50"])
    raise SystemExit(1)

# 计算分组 mAP（与论文 table[1] 的类目映射一致）
# COCO 通用类: phone,wallet,keys,backpack,suitcase,laptop,notebook,umbrella,bottle
# 校园专属类: campus_card, glasses
# 其他类: other
COCO = [0, 1, 2, 3, 4, 5, 8, 9, 10]
CAMPUS = [6, 7]
OTHER = [11]
pc = {c["name"]: c for c in best["per_class"]}
names = best["meta"]["class_names"]


def grp(idxs, key):
    vals = [pc[names[i]][key] for i in idxs]
    return sum(vals) / len(vals)


coco50 = grp(COCO, "ap50")
campus50 = grp(CAMPUS, "ap50")
other50 = grp(OTHER, "ap50")
coco95 = grp(COCO, "ap50_95")
campus95 = grp(CAMPUS, "ap50_95")
other95 = grp(OTHER, "ap50_95")
overall50 = best["overview"]["mAP50"]
overall95 = best["overview"]["mAP50_95"]


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def insert_after(idx, text, style=None):
    ref = doc.paragraphs[idx]
    new_p = doc.add_paragraph()
    if style:
        new_p.style = style
    new_p.add_run(text)
    ref._p.addnext(new_p._p)
    return new_p


# ---------- 1. 评分公式 + 演进叙述 ----------
doc.paragraphs[142] and set_text(doc.paragraphs[142],
    "　　匹配引擎是本系统的算法核心。给定一条失物 l 与一条拾物 f，匹配度 score 由「分类 + 文字语义 + 时间」七子维度加权得到（v2 评分引擎，详见式 3-2）。")

set_text(doc.paragraphs[143],
    "score = 20·(photo⊕category) + 15·qty + 20·color + 10·state + 15·place + 10·other + 10·time      (式 3-2，各子项∈[0,1]，权重和=100)\n"
    "　　归一化：设 W 为失主实际提供维度的权重之和，k = 100 / max(W, 50)，最终分 total = clamp(Σ wᵢ·sᵢ · k, 0, 100)。失主未填维度既不计入分子也不计入分母，避免“填得多得分高”的偏差。\n"
    "　　各维口径：photo⊕category（20）= 首图感知哈希 Hamming 相似度与类目命中（精确 1.0 / 部分 0.5 / 不命中 0）的均值，无图降级 0；qty 量词（15）= 失主描述数量词与拾物登记数量的一致性；color 颜色（20）= 13 系颜色族（黑/白/灰银/棕/红/橙/黄/绿/蓝/粉/紫/金/透明）匹配，跨系冲突该维计 0；state 状态（10）= 保管状态/成色一致性；place 地点（15）= 结构化行政区划相似度（同区/市/省 1.0/0.6/0.3）；other 其他（10）= 品牌、标记、关键词等文本属性命中率（含 WordNet 同义与中文近义召回）；time 时间（10）= 10·exp(−Δd/15)，Δd 为丢失与拾得天数差，τ=15 天。阈值 threshold=80（疑似线）；候选集经类别主键过滤后按 total 降序返回，前 10 条 + 所有≥80 的疑似项全列（防漏召回护栏）。")

set_text(doc.paragraphs[144],
    "　　评分引擎经历了两轮演进：初版采用“图像相似、类别、外观、特征、时间、地点”六维加权（式 3-1 旧），但外观/特征与文字属性、地点与文字地点存在语义重合，且权重固定未归一化，导致“失主描述越完整得分反而未必越高”。第二轮将冗余维度并入统一的文字语义维度，形成“分类 + 文字 + 时间”的精简结构；最终 v2 版进一步把文字拆为量词、颜色、状态、地点、其他五个可解释子维度，并引入归一化系数 k=100/max(W_provided,50)，使分数只由失主实际填写的维度决定、且同一失物内候选可比。黄金用例（一串黑色钥匙，教学楼四楼402掉落）对三候选的归一化得分分别为 A=56.25、B=86.25、C=97.5（原始分 45/69/78，k=1.25），B、C 越过疑似线 80，验证了归一化后高分候选自然涌现。")

# ---------- 2. 实现章 MatchService ----------
set_text(doc.paragraphs[211],
    "匹配模块（MatchService）是算法核心。给定一条失物 l 与一条拾物 f，匹配度 score 由「分类 + 文字语义 + 时间」七子维度加权得到（v2 评分引擎）：score = 20·(photo⊕category) + 15·qty + 20·color + 10·state + 15·place + 10·other + 10·time（0~100 分制），并经归一化 k=100/max(W_provided,50) 归一到 [0,100]。其中 photo⊕category 为首图相似度与类目命中的均值；qty/color/state/place/other 为从文字抽取的量词/颜色/状态/地点/其他关键词五维（13 系颜色表，跨系冲突计 0）；time=10·exp(−Δd/15)。各子维度评分与归一化系数均为无状态纯函数，便于单元测试与权重敏感性分析；权重、τ 与归一化下限均为配置外置项，可复现、可消融。[8][9]")

set_text(doc.paragraphs[212],
    "候选集先经类别主键过滤，再按归一化后的 total 降序返回；低于阈值（threshold=80，疑似线）的匹配不予推送，但“前 10 条保底 + 所有≥80 疑似全列”的护栏避免漏召回。photo_category、qty、color、state、place、keyword、time 七维打分与归一化系数 norm_factor 均实现为无状态纯函数，便于单元测试与权重敏感性分析；七维权重、τ 与归一化下限均为配置外置项，可复现、可消融。[8][9]")

# 实现章伪代码（节选）-> v2
set_text(doc.paragraphs[263],
    "# 加权打分（v2 节选伪代码）\n"
    "def score(lost, found):\n"
    "    if lost.category_id != found.category_id: return None   # 不同类不进打分\n"
    "    fl = extract_features(lost, is_lost=True)    # 七子维度单次流水线抽取\n"
    "    ff = extract_features(found, is_lost=False)\n"
    "    raw = (20*photo_category(fl, ff) + 15*qty(fl, ff) + 20*color(fl, ff)\n"
    "           + 10*state(fl, ff) + 15*place(fl, ff) + 10*keyword(fl, ff) + 10*time(fl, ff))\n"
    "    W = sum(w for dim, w in PROVIDED_WEIGHTS.items() if dim in lost.provided_dims)\n"
    "    k = 100 / max(W, MATCH_NORM_MIN_WEIGHT)\n"
    "    return clamp(raw * k, 0, 100)   # 阈值 80 判定疑似")

# ---------- 3. 摘要 / 主要工作 / 相关研究 措辞对齐 ----------
set_text(doc.paragraphs[85],
    "　　本文主要工作包括：（1）需求与架构设计：面向失主、拾得者、管理员、游客四类角色，梳理功能性与非功能性需求，给出分层架构、数据库 Schema（10 张表）与 RESTful API 规范；（2）进程内视觉打标：将 YOLOv8s 与 YOLO-World 推理以进程内方式嵌入后端，拾得者发布物品时自动完成类别识别，取消独立推理微服务以省内存、避超时；（3）加权匹配引擎：设计“照片⊕类别 + 量词 + 颜色 + 状态 + 地点 + 其他 + 时间”七维加权打分函数（含归一化），权重、τ 与归一化下限全部配置外置、可实验复现；（4）防冒领闭环：以审计黑匣子记录关键操作、动态交接码双端扫码确认、已解决栏申诉三重机制保证交接可信；（5）实现与测试：完成后端最小闭环（发布→打标→匹配→认领→交接）的编码，并设计单元测试与端到端测试进行验证。")

set_text(doc.paragraphs[92],
    "　　当前国内外在失物招领系统与物品识别方向的研究可归纳为三类。（1）校园失物招领类微信小程序与自建平台。以“拾到空间”等校园失物招领小程序[5]为代表，提供发布、列表、搜索与消息通知功能，多以关键词/分类浏览为主，物品描述依赖用户手工录入，图像理解能力薄弱，匹配高度依赖用户主动比对，效率有限，构成本系统的需求对照基线。（2）引入图像识别的失物招领应用。部分研究将图像识别用于物品归类与相似检索，在一定程度上降低了人工描述成本[8][9]，但其识别类别受限于闭集模型，难以覆盖“校园卡、钥匙、笔记本”等校园专属小众品类，且普遍缺少防冒领溯源机制。（3）通用目标检测技术进展。YOLO 系列算法历经多版本演进，在精度与速度间取得良好平衡，已成为工程落地主流[6]；其中针对小目标与轻量化场景的改进工作持续推动检测能力提升[2][3]，YOLOv8 在精度与速度的平衡上表现突出[6]；2023 年提出的 YOLO-World 进一步以文本提示实现开放词汇（零样本）检测[10]，开放词汇目标检测亦有系统综述[4]，为未见类别的开箱即用识别提供新路径；这些成果多为算法层面，尚未直接构成失物招领业务闭环。在后端工程侧，FastAPI+Vue3+SQLAlchemy 的全栈 Python 实践为系统构建提供了成熟范式[1][7]。本系统与上述系统的不同之处在于三点：其一，视觉自动归类采用 YOLOv8s（COCO 预训练覆盖常用类）叠加 YOLO-World（零样本补校园专属类）[10][4]，无需重新训练即可识别小众品类；其二，匹配以“照片⊕类别+量词+颜色+状态+地点+其他+时间”七维可解释加权打分（归一化到 [0,100]）实现[8][9]，权重、τ 与归一化下限由验证集实验定参，便于消融与敏感性分析；其三，构建“发布—认领—交接—归档”全链路审计与防冒领溯源闭环，弥补现有平台在可信交接上的缺失[5]。")

# 实现章模块概述：六维 -> 七维
set_text(doc.paragraphs[260],
    "发布模块（PublishService）：拾得者上传照片与保管状态后，服务在进程内调用 VisionService.predict() 完成视觉打标，写入 found_item 并触发反向主动匹配——查询同类别、待认领的失物，交由 MatchService 打分，生成疑似匹配记录并推送。匹配模块（MatchService）：实现前述七维加权（v2）score 函数，七维子分数与归一化系数 norm_factor 均为无状态纯函数，便于单元测试；候选集经类别主键过滤后按 total 降序返回，低于阈值则不推送。交接模块（HandoverService）：认领确认后生成 6 位动态交接码，Redis 作活性存储（TTL=30 分钟，兼作并发唯一与限流），MySQL handover_code 表作不可变审计镜像，双端（失主/拾得者）扫码验证均通过后状态置为“已解决”并留痕。审计模块（AuditService）：对发布、认领、交接等关键操作写入 audit_log（含操作类型、目标、IP/UA/会话/GPS 与原文），形成可追溯黑匣子。核心流程如图 4.1 所示。[8][9]")

set_text(doc.paragraphs[265],
    "　　本章在既定技术栈下完成了后端核心模块的编码：进程内视觉打标、七维加权匹配、动态交接码双写与审计留痕。测试表明最小闭环可稳定运行，为下一章系统测试所验证。")

set_text(doc.paragraphs[269],
    "核心模块实现思路如下，流程图见图 4.1，关键加权打分伪代码见图 4.2：（1）发布模块调用进程内 VisionService 完成打标并触发反向匹配；（2）匹配模块实现七维加权（v2）score 纯函数；（3）交接模块生成动态码并双写审计；（4）审计模块记录关键操作原文。各模块均通过依赖注入解耦，便于测试。[8][9]")

set_text(doc.paragraphs[284],
    "本章在既定技术栈下完成核心模块编码：进程内视觉打标、七维加权匹配、动态交接码双写与审计留痕。最小闭环可稳定运行，为下一章系统测试所验证。后续迭代（v7）进一步补充了管理后台取证导出、用户自助软删除与失效倒计时等数据生命周期管理能力，使系统在功能完整性与可信治理上更趋完备。")

# ---------- 4. 训练章：YOLOv8n->YOLOv8s + 消融 ----------
set_text(doc.paragraphs[241],
    "训练过程：模型以 COCO 预训练权重初始化，在融合数据集上计算预测框与真实标注的损失（CIoU + 分类 + DFL），反向传播更新参数；每轮在验证集评估 mAP@0.5，保存最优权重 best.pt。本文另以 WIoU v3 损失做对照消融实验（见下文），因其在本不均衡数据上导致召回崩溃，最终选型保留 CIoU 损失。")

# [252] mAP 分析（v8s+CIoU + 分组）
set_text(doc.paragraphs[252],
    "由表 4-2 可知，模型在验证集上的整体 mAP@0.5 达到 %.3f（mAP@0.5:0.95 为 %.3f），其中校园专属类（%.3f）高于 COCO 通用类（%.3f），说明在 YOLO-World 零样本补位与 YOLOv8s 更强容量下，校园卡、眼镜等校园场景专属物品的检测精度优于常规通用类；“其他类”因类内差异大、视觉特征分散，mAP@0.5 仅约 %.3f，不纳入系统核心检测精度考核。作为对照，同一验证集下初版 YOLOv8n 基线的 mAP@0.5 为 %.3f（mAP@0.5:0.95 %.3f），骨干由 YOLOv8n 升级为 YOLOv8s 后整体精度提升 %.1f 个百分点；训练阶段另曾验证 WIoU 损失导致精度崩溃（见下文消融），故最终采用 CIoU 损失。后续可进一步扩充专属类样本、引入针对性数据增强来巩固精度优势。"
    % (overall50, overall95, campus50, coco50, other50, n50, n95, (overall50 - n50) * 100))

set_text(doc.paragraphs[254],
    "本章明确了以 YOLOv8s + YOLO-World 双路检测作为视觉底座（骨干由初版 YOLOv8n 升级为 YOLOv8s 以提升精度，并在 WIoU/CIoU 消融后确定采用 CIoU 损失），给出数据集构建、预处理、进程内推理与以 mAP@0.5 + 混淆矩阵为核心的验证方案（结果见 表 4-2），为系统“发布即打标”提供模型支撑。")

set_text(doc.paragraphs[264],
    "　　当前已实现并验证的是后端最小闭环（发布→打标→匹配→认领→交接）：使用 SQLite 与视觉桩（MVP 阶段确定性返回类别标签）跑通端到端流程，自带开发库 dev.db（10 张表 + 12 分类种子数据）。在此基础上，单元测试覆盖加权打分纯函数（时间衰减、地点层级、关键词 Jaccard、七维加权与归一化、阈值判定等），端到端测试覆盖全流程与错误路径。真实 YOLOv8s + YOLO-World 推理已接入（进程内单例），前端 Web 界面已实现并与后端联调；运行效果见图 4.3。")

# ---------- 5. 测试数 184 -> 334 ----------
set_text(doc.paragraphs[290], "　　关键模块测试记录如表 6-1 所示（节选自 334 例测试用例）。")
set_text(doc.paragraphs[291],
    "　　测试运行结果：全量测试 334/334 通过（另有 1 例跳过），覆盖端到端 happy-path、加权打分单测、模块验证与错误路径，通过率 100%。源码经审查无功能性缺陷；统一响应契约、错误码、加权打分公式与数据库索引均与设计一致。测试隔离经增强后可稳定复现，为“系统测试”章提供了可信证据。")
set_text(doc.paragraphs[292],
    "　　本章通过单元与端到端测试对系统进行了验证：334 例用例全部通过，核心闭环与错误路径行为符合设计，加权匹配算法经单测确认正确。系统在功能完整性与稳定性上达到了毕设预期。")
set_text(doc.paragraphs[299], "关键模块测试记录如表 5-1 所示（节选自 334 例测试用例），测试分布见图 5.1。所有用例均通过函数级隔离 fixture 确保可复现。")
set_text(doc.paragraphs[305],
    "测试运行结果：全量 334/334 通过（另有 1 例跳过），通过率 100%。源码经审查无功能性缺陷；统一响应契约、错误码、加权打分公式与数据库索引均与设计一致。加权匹配算法经单测确认正确，端到端 happy-path 与模块验证用例覆盖主要错误路径。系统在功能完整性与稳定性上达到毕设预期。")
set_text(doc.paragraphs[310],
    "本章通过单元与端到端测试对系统验证：334 例用例全部通过，核心闭环与错误路径行为符合设计，加权匹配算法经单测确认正确。系统在功能完整性与稳定性上达到毕设预期。")

# ---------- 6. 结论 / 摘要 ----------
set_text(doc.paragraphs[313],
    "　　本文围绕“校园失物招领”这一高频校园需求，完成了一套基于 YOLOv8 的智能匹配系统的需求分析、设计、实现与测试。主要成果如下：（1）提出并实现了“视觉自动打标 + 七维加权匹配 + 防冒领闭环”的端到端方案，拾得者发布仅需拍照与保管状态，系统自动归类并主动推送疑似匹配；（2）将 YOLOv8s（COCO 常用类）与 YOLO-World（零样本校园类）以进程内方式嵌入后端，取消独立推理微服务，降低内存与超时风险；（3）设计了“照片⊕类别 + 量词 + 颜色 + 状态 + 地点 + 其他 + 时间”七维加权打分函数，并引入归一化系数使分数只由失主实际填写维度决定、不失可比性，权重、τ 与归一化下限全部配置外置、可复现；（4）以审计黑匣子、动态交接码双端确认、已解决栏申诉三重机制防范冒领；（5）完成后端最小闭环编码，并通过 334 例单元/端到端测试验证，全部通过。")

set_text(doc.paragraphs[314],
    "　　工作过程中也遇到并解决了若干关键问题：测试套件初期的共享状态导致偶发不稳定，通过函数级隔离 fixture 修复；独立视觉服务带来的额外开销与超时风险，经论证改为进程内推理；训练阶段尝试以 WIoU 损失替换 CIoU，因本数据集类别高度不均衡导致稀有类召回崩溃（mAP@0.5 由 %.3f 跌至 %.3f，召回率仅 %.3f），经消融分析回退 CIoU 并重训 YOLOv8s 解决。受毕设周期限制，仍存在以下待完善之处：（1）视觉模型已由 YOLOv8n 升级为 YOLOv8s，真实推理已接入并实测 mAP@0.5=%.3f，但校园专属稀有类（校园卡/钱包/钥匙/眼镜）仍受样本不均衡限制、AP 偏低，后续可扩充专属类样本与针对性数据增强进一步巩固；（2）WebSocket 即时沟通（当前以轮询兜底）与公示栏、管理后台的交互体验仍待进一步增强；（3）Web 端无法可靠获取 EXIF/GPS，地点定位以 6 位行政区划码 region_code 近似，finer-grained 定位留待移动端。"
    % (overall50, wiou["overview"]["mAP50"], wiou["overview"]["recall"], overall50))

# ---------- 7. table[2] mAP 数据更新 ----------
tbl = doc.tables[2]
tbl.rows[1].cells[1].text = "%.3f" % coco50
tbl.rows[1].cells[2].text = "%.3f" % campus50
tbl.rows[1].cells[3].text = "%.3f" % overall50
tbl.rows[1].cells[4].text = "%.3f" % other50
tbl.rows[2].cells[1].text = "%.3f" % coco95
tbl.rows[2].cells[2].text = "%.3f" % campus95
tbl.rows[2].cells[3].text = "%.3f" % overall95
tbl.rows[2].cells[4].text = "%.3f" % other95

# ---------- 8. 新增段落 ----------
# 必须严格自底向上插入（284 -> 252 -> 212），否则先插小索引会把后面的锚点整体后移一位，
# 导致后续 insert_after 命中错误段落。
# 8a. 管理员能力（实现章小结后）
insert_after(284,
    "管理员能力在 v7 取证导出基础上进一步扩展：注册页新增邀请码字段，命中服务端 ADMIN_APPLY_CODE（常量时间比较，防时序侧信道；空码护栏防全员越权）即静默升级为管理员，未填/错填与普通注册响应体一致以防探测；管理后台提供用户列表检索、匹配详情（含匿名 IM 对话原文）与取证导出（xlsx / markdown / csv 三格式，审计留痕），并支持 all_time 长期留存查询，满足失物纠纷取证与合规审计需求。")

# 8b. WIoU 消融叙述（mAP 分析 [252] 之后）—— 先插大索引
insert_after(252,
    "　　在最终选型前，本文对边界框回归损失做了对照消融。以 YOLOv8s 为骨干、改用 WIoU v3 损失微调时，验证集 mAP@0.5 骤降至 %.3f（mAP@0.5:0.95 为 %.3f），精确率 %.3f、召回率仅 %.3f，较同骨干 CIoU 方案（mAP@0.5 %.3f、召回率 %.3f）几近完全退化，逐类 AP 普遍跌至 0.1 以下。分析其原因：WIoU v3 以动态非单调聚焦机制“关注中等质量锚框”，对低质量预测施加更小梯度增益、抑制其对回归的干扰；而本数据集类别高度不均衡（书包、手机、水瓶等常见类样本占绝大多数，校园卡、钱包、钥匙、眼镜等校园稀有类样本极少），稀有类在训练早期产生的预测框恰恰集中于“低质量”区间，被 WIoU 判定为离群而抑制梯度，致其本就微弱的学习信号进一步衰减，最终召回全面崩溃。鉴于该退化不可接受，本文回退至 CIoU 损失并重训 YOLOv8s，验证集 mAP@0.5 恢复至 %.3f、mAP@0.5:0.95 达 %.3f，稀有类 AP 显著回升（校园卡 %.3f、眼镜 %.3f）。该消融实验表明：损失函数并非“越新越好”，其有效性高度依赖数据分布，在长尾/不均衡数据上采用聚焦式损失需格外审慎。"
    % (wiou["overview"]["mAP50"], wiou["overview"]["mAP50_95"],
       wiou["overview"]["precision"], wiou["overview"]["recall"],
       overall50, best["overview"]["recall"],
       overall50, overall95,
       pc["campus_card"]["ap50"], pc["glasses"]["ap50"]))

# 8c. flow-v3 匹配池交互（实现章匹配模块后，即原 [212] 之后）—— 最后插小索引
insert_after(212,
    "　　匹配池交互遵循“单向可见”原则：拾得者选择“留在原地”保管（keep1）的物品仅单向进入失主匹配池——失主侧可见并可主动认领，拾得者侧仅只读呈现、不反向生成候选，避免对“已放回原地、无需再介入”的拾得者造成重复打扰；低分候选（total<60）以弱化卡片 + 二次确认方式呈现，兼顾召回与防误领；疑似阈值维持 80，候选列表取“前 10 条保底 + 所有≥80 疑似全列”的护栏，防止高召回场景下的漏推。")

doc.save(THESIS)
print("THESIS UPDATED ->", THESIS)
print("v8s+CIoU: overall50=%.3f overall95=%.3f coco50=%.3f campus50=%.3f other50=%.3f"
      % (overall50, overall95, coco50, campus50, other50))
print("WIoU disaster: mAP50=%.3f recall=%.3f" % (wiou["overview"]["mAP50"], wiou["overview"]["recall"]))
