import traceback
try:
    from docx import Document
except ImportError:
    print("PYTHON_DOCX_MISSING")
    raise SystemExit(0)

PATH = r"D:/Zhuomian/毕业论文/2026年毕业设计论文模板/曹灏天计算机学院毕业论文-2026版（7-6）.docx"
doc = Document(PATH)
print("=== PARAGRAPHS (keyword hits) ===")
kw = ["评分", "加权", "mAP", "维度", "匹配度", "归一化", "WIoU", "六维", "五维", "四维",
      "类别", "关键词", "相似度", "match", "score", "Score", "加权打分"]
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    if any(k in t for k in kw):
        print(f"[{i}] {t[:260]}")
print("=== TABLES ===")
for ti, tbl in enumerate(doc.tables):
    print(f"--- table[{ti}] rows={len(tbl.rows)} cols={len(tbl.columns)} ---")
    for r in tbl.rows:
        cells = [c.text.strip() for c in r.cells]
        print(" | ".join(cells))
